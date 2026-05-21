"""Document generation HTTP routes — Download, Preview, Print."""
from __future__ import annotations
from pathlib import Path

from flask import Blueprint, request, send_file

from ..services import doc_generator
from ..database import fetch_one, fetch_all
from ..utils import envelope_error, envelope_ok, get_json_body

bp = Blueprint("document", __name__, url_prefix="/api")


@bp.get("/document/kinds")
def document_kinds():
    return envelope_ok({
        "kinds": doc_generator.VALID_KINDS,
        "templates": {
            k: str(doc_generator.config.TEMPLATES_DIR / v)
            for k, v in doc_generator.KIND_TO_TEMPLATE.items()
        },
    })


@bp.post("/cases/<case_id>/document/<kind>")
def generate_document(case_id: str, kind: str):
    """Generate a document. Optional JSON body with extra placeholder values."""
    body = get_json_body(request) or {}
    extra = body.get("extra", body)
    user = body.get("user", "system")
    try:
        result = doc_generator.generate(case_id, kind, extra=extra, user=user)
    except (ValueError, NotImplementedError) as e:
        return envelope_error(str(e), status=400, code="BAD_REQUEST")
    except LookupError as e:
        return envelope_error(str(e), status=404, code="NOT_FOUND")
    return envelope_ok(result)


@bp.get("/documents/<int:doc_id>")
def fetch_document(doc_id: int):
    """Download a previously-generated document by row id."""
    row = fetch_one("SELECT * FROM documents WHERE id=?", (doc_id,))
    if not row:
        return envelope_error("Document not found", status=404,
                              code="NOT_FOUND")
    p = Path(row["file_path"])
    if not p.exists():
        return envelope_error("File missing on disk", status=410,
                              code="FILE_MISSING")
    return send_file(str(p), as_attachment=True,
                     download_name=row["document_name"],
                     mimetype=row["mime_type"])


# ===================================================================
# PREVIEW — Show document in browser (HTML render) with Print button
# ===================================================================
@bp.get("/documents/<int:doc_id>/preview")
def preview_document(doc_id: int):
    """
    Render a .docx document as HTML for in-browser preview + print.
    No download needed — directly print from browser.
    """
    row = fetch_one("SELECT * FROM documents WHERE id=?", (doc_id,))
    if not row:
        return "<h1>Document not found</h1>", 404
    p = Path(row["file_path"])
    if not p.exists():
        return "<h1>File missing on disk</h1>", 410

    # Convert .docx to HTML
    html_content = _docx_to_html(p)

    # Wrap in a printable page with controls
    page = f"""<!DOCTYPE html>
<html><head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{row['document_name']} — Preview</title>
<style>
@media print {{
    .no-print {{ display: none !important; }}
    body {{ margin: 0; padding: 0; }}
    .doc-content {{ border: none !important; box-shadow: none !important; padding: 20mm !important; }}
}}
body {{ font-family: 'Noto Sans Devanagari', Arial, sans-serif; background: #e0e0e0; margin: 0; padding: 0; }}
.toolbar {{ position: sticky; top: 0; z-index: 100; background: linear-gradient(135deg, #667eea, #764ba2);
            padding: 12px 20px; display: flex; align-items: center; gap: 12px; box-shadow: 0 2px 10px rgba(0,0,0,0.3); }}
.toolbar h3 {{ color: #fff; margin: 0; font-size: 16px; flex: 1; }}
.toolbar .btn {{ padding: 8px 18px; border: none; border-radius: 6px; font-size: 14px;
                 font-weight: 700; cursor: pointer; text-decoration: none; }}
.btn-print {{ background: #4caf50; color: #fff; }}
.btn-download {{ background: #2196f3; color: #fff; }}
.btn-back {{ background: rgba(255,255,255,0.2); color: #fff; }}
.doc-content {{ max-width: 210mm; margin: 20px auto; background: #fff; padding: 25mm;
                box-shadow: 0 4px 20px rgba(0,0,0,0.15); border-radius: 4px;
                min-height: 297mm; line-height: 1.6; }}
.doc-content table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
.doc-content td, .doc-content th {{ border: 1px solid #333; padding: 6px 10px; font-size: 13px; }}
.doc-content th {{ background: #f0f0f0; font-weight: bold; }}
.doc-content h1 {{ font-size: 20px; text-align: center; margin-bottom: 10px; }}
.doc-content h2 {{ font-size: 16px; text-align: center; color: #555; margin-bottom: 15px; }}
.doc-content h3 {{ font-size: 14px; margin-top: 15px; border-bottom: 1px solid #ccc; padding-bottom: 4px; }}
.doc-content p {{ margin: 6px 0; font-size: 13px; }}
</style>
</head><body>

<div class="toolbar no-print">
    <h3>{row['document_name']}</h3>
    <button class="btn btn-print" onclick="window.print()">🖨️ Print</button>
    <a class="btn btn-download" href="/api/documents/{doc_id}">⬇️ Download</a>
    <button class="btn btn-back" onclick="history.back()">← Back</button>
</div>

<div class="doc-content">
{html_content}
</div>

</body></html>"""
    return page


# ===================================================================
# PREVIEW for case — generate + preview in one step
# ===================================================================
@bp.get("/cases/<case_id>/preview/<kind>")
def preview_case_document(case_id: str, kind: str):
    """
    Generate a document for a case and show preview immediately.
    If already generated, shows the latest one.
    """
    # Check if already generated recently
    existing = fetch_one(
        """SELECT id FROM documents
           WHERE case_id=? AND document_type=? AND status='active'
           ORDER BY created_at DESC LIMIT 1""",
        (case_id, kind),
    )

    if existing:
        doc_id = existing["id"]
    else:
        # Generate new
        try:
            result = doc_generator.generate(case_id, kind, user="preview")
            doc_id = result["id"]
        except (ValueError, NotImplementedError) as e:
            return f"<h1>Error: {e}</h1>", 400
        except LookupError as e:
            return f"<h1>Case not found: {e}</h1>", 404

    # Redirect to preview
    from flask import redirect
    return redirect(f"/api/documents/{doc_id}/preview")


# ===================================================================
# LIST all documents for a case with preview/download/print links
# ===================================================================
@bp.get("/cases/<case_id>/documents/page")
def case_documents_page(case_id: str):
    """HTML page listing all documents for a case with action buttons."""
    docs = fetch_all(
        """SELECT id, document_type, document_name, file_size, created_at
           FROM documents WHERE case_id=? AND status='active'
           ORDER BY created_at DESC""",
        (case_id,),
    )

    rows_html = ""
    for d in docs:
        size_kb = (d["file_size"] or 0) / 1024
        rows_html += f"""
        <tr>
            <td>{d['document_type']}</td>
            <td>{d['document_name']}</td>
            <td>{size_kb:.1f} KB</td>
            <td>{d['created_at'] or ''}</td>
            <td>
                <a href="/api/documents/{d['id']}/preview" class="btn btn-sm btn-view">👁 Preview</a>
                <a href="/api/documents/{d['id']}" class="btn btn-sm btn-dl">⬇ Download</a>
            </td>
        </tr>"""

    # Generate buttons for each kind
    gen_buttons = ""
    for kind in doc_generator.VALID_KINDS:
        gen_buttons += f'<a href="/api/cases/{case_id}/preview/{kind}" class="btn btn-gen">{kind}</a> '

    return f"""<!DOCTYPE html>
<html><head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Documents — {case_id}</title>
<style>
body {{ font-family: Arial, sans-serif; background: #f5f5f5; padding: 20px; }}
h1 {{ color: #333; }}
table {{ width: 100%; border-collapse: collapse; background: #fff; border-radius: 8px; overflow: hidden; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }}
th, td {{ padding: 10px 14px; text-align: left; border-bottom: 1px solid #eee; font-size: 14px; }}
th {{ background: #667eea; color: #fff; }}
.btn {{ display: inline-block; padding: 5px 12px; border-radius: 5px; text-decoration: none; font-size: 12px; font-weight: 600; margin: 2px; }}
.btn-sm {{ font-size: 11px; }}
.btn-view {{ background: #4caf50; color: #fff; }}
.btn-dl {{ background: #2196f3; color: #fff; }}
.btn-gen {{ background: #ff9800; color: #fff; padding: 8px 14px; margin: 4px; border-radius: 6px; font-size: 13px; }}
.section {{ background: #fff; padding: 16px; border-radius: 8px; margin-bottom: 16px; box-shadow: 0 2px 6px rgba(0,0,0,0.05); }}
</style>
</head><body>
<h1>Documents — {case_id}</h1>

<div class="section">
<h3>Generate New Document:</h3>
{gen_buttons}
</div>

<div class="section">
<h3>Existing Documents ({len(docs)}):</h3>
<table>
<tr><th>Type</th><th>File</th><th>Size</th><th>Created</th><th>Actions</th></tr>
{rows_html if rows_html else '<tr><td colspan="5" style="text-align:center;color:#999;">No documents yet</td></tr>'}
</table>
</div>

</body></html>"""


@bp.post("/templates/migrate-legacy")
def migrate_legacy():
    """Convert «FIELD» placeholders in a template file to {{ FIELD }}."""
    body = get_json_body(request) or {}
    file_name = body.get("file")
    if not file_name:
        return envelope_error("Provide JSON {\"file\": \"<template.docx>\"}",
                              status=400)
    src = doc_generator.config.TEMPLATES_DIR / file_name
    if not src.exists():
        return envelope_error(f"Template not found: {src}", status=404,
                              code="NOT_FOUND")
    backup = doc_generator.migrate_legacy_template_in_place(src)
    return envelope_ok({"migrated": str(src), "backup": str(backup)})


# ===================================================================
# Helper: Convert .docx to HTML for preview
# ===================================================================
def _docx_to_html(docx_path: Path) -> str:
    """Convert a .docx file to basic HTML for browser preview."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        doc = Document(str(docx_path))
    except Exception as e:
        return f"<p style='color:red'>Error reading document: {e}</p>"

    html_parts = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragraph
            from docx.oxml.ns import qn
            pPr = element.find(qn("w:pPr"))
            style_name = ""
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None:
                    style_name = pStyle.get(qn("w:val"), "")

            # Extract text with bold/italic
            text_parts = []
            for run in element.findall(f".//{{{element.nsmap.get('w', '')}}}r"):
                rPr = run.find(f"{{{element.nsmap.get('w', '')}}}rPr")
                t = run.find(f"{{{element.nsmap.get('w', '')}}}t")
                txt = t.text if t is not None and t.text else ""

                is_bold = False
                is_italic = False
                if rPr is not None:
                    if rPr.find(f"{{{element.nsmap.get('w', '')}}}b") is not None:
                        is_bold = True
                    if rPr.find(f"{{{element.nsmap.get('w', '')}}}i") is not None:
                        is_italic = True

                if is_bold and is_italic:
                    text_parts.append(f"<b><i>{txt}</i></b>")
                elif is_bold:
                    text_parts.append(f"<b>{txt}</b>")
                elif is_italic:
                    text_parts.append(f"<i>{txt}</i>")
                else:
                    text_parts.append(txt)

            full_text = "".join(text_parts)

            if "Heading" in style_name:
                level = "".join(c for c in style_name if c.isdigit()) or "3"
                level = min(int(level) + 1, 4)  # Heading0→h1, Heading1→h2, etc
                html_parts.append(f"<h{level}>{full_text}</h{level}>")
            elif "Title" in style_name:
                html_parts.append(f"<h1>{full_text}</h1>")
            else:
                if full_text.strip():
                    html_parts.append(f"<p>{full_text}</p>")
                else:
                    html_parts.append("<p>&nbsp;</p>")

        elif tag == "tbl":
            # Table
            html_parts.append("<table>")
            for row_el in element.findall(f".//{{{element.nsmap.get('w', '')}}}tr"):
                html_parts.append("<tr>")
                for cell_el in row_el.findall(f".//{{{element.nsmap.get('w', '')}}}tc"):
                    cell_text_parts = []
                    for p in cell_el.findall(f".//{{{element.nsmap.get('w', '')}}}p"):
                        for t in p.findall(f".//{{{element.nsmap.get('w', '')}}}t"):
                            if t.text:
                                cell_text_parts.append(t.text)
                    cell_text = " ".join(cell_text_parts)
                    html_parts.append(f"<td>{cell_text}</td>")
                html_parts.append("</tr>")
            html_parts.append("</table>")

    return "\n".join(html_parts)



# ===================================================================
# COMBINE ALL — Merge all case documents into one file for printing
# ===================================================================
@bp.get("/cases/<case_id>/documents/combine")
def combine_all_documents(case_id: str):
    """
    Combine ALL documents of a case into a single merged .docx file.
    Download as one file → print everything at once.

    Query params:
      ?types=provisional_consumer,section3,noc  (optional filter)
      If no types specified, combines ALL documents.
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime
    import io

    docs = fetch_all(
        """SELECT * FROM documents WHERE case_id=? AND status='active'
           ORDER BY created_at ASC""",
        (case_id,),
    )
    if not docs:
        return envelope_error("No documents found for this case", status=404)

    # Optional filter by types
    type_filter = request.args.get("types")
    if type_filter:
        allowed = set(t.strip() for t in type_filter.split(","))
        docs = [d for d in docs if d["document_type"] in allowed]

    if not docs:
        return envelope_error("No matching documents", status=404)

    # Create combined document
    combined = Document()

    # Cover page
    h = combined.add_heading(f"Case File: {case_id}", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    combined.add_paragraph(f"Generated: {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    combined.add_paragraph(f"Total Documents: {len(docs)}")
    combined.add_paragraph("")

    # Table of contents
    combined.add_heading("Index / सूची", level=1)
    toc_table = combined.add_table(rows=1, cols=4)
    toc_table.style = "Table Grid"
    hdr = toc_table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Document Type"
    hdr[2].text = "File Name"
    hdr[3].text = "Date"
    for i, d in enumerate(docs, 1):
        row = toc_table.add_row().cells
        row[0].text = str(i)
        row[1].text = d["document_type"] or ""
        row[2].text = d["document_name"] or ""
        row[3].text = (d["created_at"] or "")[:10]

    # Add each document
    for i, doc_row in enumerate(docs):
        # Page break before each document
        combined.add_page_break()

        # Section header
        combined.add_heading(
            f"[{i+1}/{len(docs)}] {doc_row['document_type']} — {doc_row['document_name']}",
            level=1)
        combined.add_paragraph(f"Created: {doc_row['created_at'] or 'N/A'}")
        combined.add_paragraph("")

        # Try to merge .docx content
        fpath = Path(doc_row["file_path"])
        if fpath.exists() and fpath.suffix.lower() == ".docx":
            try:
                sub_doc = Document(str(fpath))
                for element in sub_doc.element.body:
                    combined.element.body.append(element)
            except Exception as e:
                combined.add_paragraph(f"[Could not merge: {e}]")
        elif fpath.exists():
            combined.add_paragraph(
                f"[Non-docx file: {doc_row['document_name']} "
                f"({doc_row['mime_type']}) — see separate download]")
        else:
            combined.add_paragraph("[File missing on disk]")

    # Save to memory
    output = io.BytesIO()
    combined.save(output)
    output.seek(0)

    filename = f"COMPLETE_CASE_{case_id}_{datetime.now().strftime('%Y%m%d')}.docx"
    return send_file(
        output, as_attachment=True, download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ===================================================================
# COMBINE + PREVIEW — all docs in browser for one-shot printing
# ===================================================================
@bp.get("/cases/<case_id>/documents/combine-preview")
def combine_preview(case_id: str):
    """
    Show ALL documents of a case combined in browser for one-shot print.
    No download — just preview and print button.
    """
    docs = fetch_all(
        """SELECT * FROM documents WHERE case_id=? AND status='active'
           ORDER BY created_at ASC""",
        (case_id,),
    )

    all_html = ""
    for i, d in enumerate(docs):
        fpath = Path(d["file_path"])
        if fpath.exists() and fpath.suffix.lower() == ".docx":
            content = _docx_to_html(fpath)
        elif fpath.exists() and d.get("mime_type", "").startswith("image/"):
            content = f'<img src="/api/documents/{d["id"]}" style="max-width:100%;border:1px solid #ccc;">'
        else:
            content = f'<p><a href="/api/documents/{d["id"]}">Download: {d["document_name"]}</a></p>'

        all_html += f"""
        <div class="doc-section">
            <div class="doc-header">[{i+1}/{len(docs)}] {d['document_type']} — {d['document_name']}</div>
            {content}
        </div>
        <div class="page-break"></div>
        """

    return f"""<!DOCTYPE html>
<html><head>
<meta charset="utf-8">
<title>Complete Case File — {case_id}</title>
<style>
@media print {{
    .no-print {{ display: none !important; }}
    .page-break {{ page-break-after: always; }}
    body {{ margin: 0; }}
}}
body {{ font-family: 'Noto Sans Devanagari', Arial, sans-serif; background: #e0e0e0; }}
.toolbar {{ position: sticky; top: 0; z-index: 100; background: linear-gradient(135deg, #667eea, #764ba2);
            padding: 12px 20px; display: flex; align-items: center; gap: 12px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.3); }}
.toolbar h3 {{ color: #fff; margin: 0; flex: 1; font-size: 15px; }}
.toolbar .btn {{ padding: 8px 16px; border: none; border-radius: 6px; font-size: 13px;
                 font-weight: 700; cursor: pointer; text-decoration: none; color: #fff; }}
.btn-print {{ background: #4caf50; }}
.btn-download {{ background: #2196f3; }}
.btn-back {{ background: rgba(255,255,255,0.2); }}
.doc-section {{ max-width: 210mm; margin: 15px auto; background: #fff; padding: 20mm;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1); min-height: 200px; }}
.doc-header {{ background: #667eea; color: #fff; padding: 8px 14px; border-radius: 6px;
               font-weight: bold; margin-bottom: 15px; font-size: 13px; }}
.doc-section table {{ border-collapse: collapse; width: 100%; margin: 8px 0; }}
.doc-section td, .doc-section th {{ border: 1px solid #333; padding: 5px 8px; font-size: 12px; }}
.doc-section h1,.doc-section h2,.doc-section h3 {{ text-align: center; }}
.page-break {{ height: 5px; }}
</style>
</head><body>
<div class="toolbar no-print">
    <h3>Complete Case: {case_id} ({len(docs)} documents)</h3>
    <button class="btn btn-print" onclick="window.print()">🖨️ Print All</button>
    <a class="btn btn-download" href="/api/cases/{case_id}/documents/combine">⬇️ Download .docx</a>
    <button class="btn btn-back" onclick="history.back()">← Back</button>
</div>
{all_html if docs else '<div class="doc-section"><p style="text-align:center;color:#999;">No documents found</p></div>'}
</body></html>"""


# ===================================================================
# APPEAL UPLOAD shortcut — upload appeal application directly
# ===================================================================
@bp.post("/cases/<case_id>/appeal-upload")
def appeal_upload_shortcut(case_id: str):
    """
    Quick upload for appeal documents without creating an appeal record first.
    Creates appeal record automatically + attaches the file.

    Multipart form:
      - file: PDF/image of appeal application
      - appellant_name: (optional, default from case)
      - appeal_reason: (optional)
    """
    import uuid, os
    from datetime import datetime

    case = fetch_one("SELECT * FROM raid_cases WHERE case_id=?", (case_id,))
    if not case:
        return envelope_error("Case not found", status=404)

    if "file" not in request.files:
        return envelope_error("No file provided", status=400)

    file = request.files["file"]
    if not file or file.filename == "":
        return envelope_error("Empty file", status=400)

    ext = Path(file.filename).suffix.lower()
    if ext not in config.ALLOWED_UPLOAD_EXTENSIONS:
        return envelope_error(f"Type not allowed", status=400)

    file.seek(0, 2)
    size = file.tell()
    file.seek(0)
    if size > config.MAX_UPLOAD_SIZE_BYTES:
        return envelope_error(f"Max {config.MAX_UPLOAD_SIZE_MB} MB", status=413)

    # Auto-create appeal record
    appellant = request.form.get("appellant_name") or case.get("user_name") or "Consumer"
    reason = request.form.get("appeal_reason") or "See attached document"

    from ..database import execute as db_execute
    cur = db_execute(
        """INSERT INTO appeals
              (case_id, appeal_date, appellant_name, appellant_relation,
               appeal_reason, appeal_status)
           VALUES (?,?,?,?,?,?)""",
        (case_id, datetime.now().strftime("%Y-%m-%d"), appellant,
         request.form.get("relationship", "self"), reason, "received"))
    appeal_id = cur.lastrowid

    # Update case status
    db_execute(
        "UPDATE raid_cases SET case_status='appealed', updated_at=datetime('now') "
        "WHERE case_id=?", (case_id,))

    # Save file
    case_dir = config.UPLOADS_DIR / case_id / "appeals"
    case_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    uid = uuid.uuid4().hex[:8]
    safe_name = "".join(c if (c.isalnum() or c in "._-") else "_"
                        for c in os.path.basename(file.filename)) or "appeal"
    stored = f"appeal_{appeal_id}_{ts}_{uid}_{safe_name}"
    fp = case_dir / stored
    file.save(str(fp))
    actual_size = fp.stat().st_size

    mime_map = {".pdf": "application/pdf", ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg", ".png": "image/png",
                ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    mime = mime_map.get(ext, "application/octet-stream")

    doc_cur = db_execute(
        """INSERT INTO documents
              (case_id, document_type, document_name, file_path,
               file_size, mime_type, uploaded_by, status)
           VALUES (?,?,?,?,?,?,?,?)""",
        (case_id, f"appeal_proceeding_{appeal_id}", safe_name,
         str(fp), actual_size, mime,
         request.form.get("uploaded_by", "system"), "active"))

    return envelope_ok({
        "appeal_id": appeal_id,
        "doc_id": doc_cur.lastrowid,
        "case_status": "appealed",
        "file_name": safe_name,
        "message": "Appeal filed + document uploaded / अपील दर्ज + दस्तावेज अपलोड",
    })
