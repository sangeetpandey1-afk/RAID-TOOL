"""
Document generation engine.

Two paths:

1. **Template-driven** (preferred for production)
   * Place a `.docx` file in ``templates/`` matching the document kind.
   * The template uses Jinja-style placeholders: ``{{ NAME }}``.
   * If your existing templates use the legacy ``«NAME»`` notation, run
     :func:`migrate_legacy_template_in_place` once — it rewrites the file.

2. **Auto-generated** (fallback when no template is present)
   * The system writes a clean professional .docx from scratch using
     ``python-docx`` so the officer is never blocked.

Output files land in ``docs/<case_id>/<kind>_<timestamp>.docx`` and a row is
written to the ``documents`` table for tracking.
"""
from __future__ import annotations
import logging
import re
import shutil
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docxtpl import DocxTemplate

from .. import config
from ..database import execute, fetch_all, fetch_one
from ..utils import from_json_str, safe_float

log = logging.getLogger(__name__)


# ===================================================================
# Document kind catalogue
# ===================================================================
KIND_TO_TEMPLATE: dict[str, str] = {
    "provisional_consumer": "provisional_consumer.docx",
    "provisional_office":   "provisional_office.docx",
    "provisional_pvvnl":    "provisional_pvvnl.docx",  # exact PVVNL/UPPCL format
    "section3":             "section3_notice.docx",
    "section5":             "section5_notice.docx",
    "thanedari":            "thanedari_copy.docx",
    "envelope":             "envelope.docx",
    "deposit_slip":         "deposit_slip.docx",
    "noc":                  "noc.docx",
    "compounding_order":    "compounding_order.docx",
}

VALID_KINDS = list(KIND_TO_TEMPLATE.keys())


# ===================================================================
# Context builder
# ===================================================================
def _money(v: Any) -> str:
    f = safe_float(v)
    return f"{f:,.2f}"


def build_context(case: dict, consumer: dict | None,
                  assessment: dict | None,
                  *, extra: dict | None = None) -> dict:
    """Produce the dict of placeholders for the template (and the auto-doc)."""
    extra = extra or {}
    cons = consumer or {}
    a = assessment or {}
    # Prefer the computed device list (with L/F/H/D/units) from the
    # assessment block; fall back to the raw input devices on the case.
    devices = (a.get("devices") if a.get("devices") else case.get("devices")) or []
    fixed = a.get("fixed_charges") or {}
    energy = a.get("energy_charges") or {}
    ed = a.get("electricity_duty") or {}
    slabs = energy.get("slabs") or []

    today_iso = date.today().isoformat()

    # Load office config from system_config (cached per-call)
    cfg = _load_office_config()

    ctx: dict[str, Any] = {
        # ---- Core
        "Div_no":          cons.get("div_code") or case.get("div_no") or "",
        "ONLINE_NO":       case.get("online_no") or "",
        "disno":           case.get("case_id") or "",
        "ESIFO":           case.get("fir_number") or "",
        "CH_no":           extra.get("ch_no", ""),
        "dis_date":        case.get("inspection_date") or today_iso,
        "Date":            today_iso,
        "ASSESMENT_TOTAL": _money(case.get("total_assessment") or a.get("grand_total")),
        "date1":           today_iso,
        "time11":          datetime.now().strftime("%H:%M"),
        "revice":          extra.get("revice", ""),

        # ---- Consumer + user details
        "div_no":          cons.get("div_code") or "",
        "NAME":            cons.get("name") or "",
        "father_nane":     cons.get("father_name") or "",
        "USER_NAME":       case.get("user_name") or cons.get("name") or "",
        "USERS_FATHER":    case.get("user_father") or cons.get("father_name") or "",
        "VILLAGE":         cons.get("village") or "",
        "post":            cons.get("post_office") or "",
        "pin_code":        cons.get("pin_code") or "",
        "conaction_no":    case.get("account_number") or cons.get("account_number") or "",
        "MOBILE_NO":       cons.get("mobile") or "",
        "ACCOUNT_ID":      case.get("account_number") or cons.get("account_number") or "",
        "CONNECTION_LOAD": cons.get("load_value") or case.get("connected_load_kw") or "",
        "SUPPLY_TYPE":     cons.get("supply_type") or "",
        "JE_NAME":         case.get("je_name") or "",
        "CHECKING_TYPE":   case.get("checking_type") or "",
        "SUB_SUBSTATION":  case.get("sub_substation") or cons.get("sub_substation") or "",
        "CATEGORY":        cons.get("category") or "",
        "TD_DATE":         case.get("td_date") or "",
        "LANDMARK":        cons.get("landmark") or "",
        "tehsil":          cons.get("tehsil") or "",
        "district":        cons.get("district") or "",
        "post_office":     cons.get("post_office") or "",
        "consumer_name":   cons.get("name") or "",
        "consumer_father": cons.get("father_name") or "",
        "village":         cons.get("village") or "",
        "FIR_NUMBER":      case.get("fir_number") or "",
        "SECTION":         case.get("section") or "",
        "COMPOUNDING_AMOUNT": _money(case.get("compounding_amount")),

        # ---- Section 3 fields
        "sec3_no":         extra.get("sec3_no", ""),
        "sec3_date":       extra.get("sec3_date", today_iso),
        "sec3_amaunt":     _money(case.get("total_assessment")),
        "total_sec3":      _money((safe_float(case.get("total_assessment"))
                                   + config.ADMIN_FEE_SECTION_3)),

        # ---- Section 5 fields
        "rc_number":       extra.get("rc_number", ""),
        "letter_number":   extra.get("letter_number", ""),
        "current_date":    today_iso,
        "outstanding_amount": _money(case.get("total_assessment")),
        "checking_report_number": extra.get("checking_report_number", ""),
        "checking_date":   case.get("inspection_date") or "",
        "demand_notice_number": extra.get("demand_notice_number", ""),
        "demand_notice_date":   extra.get("demand_notice_date", ""),
        "grid_number":     extra.get("grid_number", ""),
        "grid_date":       extra.get("grid_date", ""),

        # ---- Calculation summary (consumer copy)
        "CONNECTED_LOAD":     cons.get("load_value") or case.get("connected_load_kw") or "",
        "FIXED_RATE":         _money(fixed.get("fixed_rate")),
        "MONTHS":             fixed.get("months") or a.get("months") or "",
        "FIXED_AMOUNT":       _money(fixed.get("base")),
        "FINAL_FIXED_CHARGES": _money(fixed.get("final")),
        "FINAL_ENERGY_CHARGES": _money(energy.get("final")),
        "FINAL_ED_CHARGES":   _money(ed.get("amount")),
        "TOTAL_UNITS":        a.get("total_units_after_less_unit") or a.get("total_units_calculated") or "",
        "ED_RATE_PERCENT":    ed.get("ed_percent") or "",
        "ED_AMOUNT":          _money(ed.get("amount")),

        # ---- Office-copy LFHD totals
        "TOTAL_CALCULATED_UNITS": a.get("total_units_calculated") or "",
        "SUMMARY_FIXED":   _money(fixed.get("final")),
        "SUMMARY_ENERGY":  _money(energy.get("final")),
        "SUMMARY_ED":      _money(ed.get("amount")),
        "SUMMARY_TOTAL":   _money(a.get("grand_total")),
    }

    # Devices — up to 30 slots with both DEVICE_n_* (consumer copy) and L_n etc (office)
    for i, d in enumerate(devices, start=1):
        ctx[f"DEVICE_{i}_NAME"]  = d.get("name") or ""
        ctx[f"DEVICE_{i}_LOAD"]  = d.get("L") or d.get("load") or ""
        ctx[f"DEVICE_{i}_UNITS"] = d.get("units") or ""
        ctx[f"L_{i}"] = d.get("L") or d.get("load") or ""
        ctx[f"F_{i}"] = d.get("F") or d.get("factor") or ""
        ctx[f"H_{i}"] = d.get("H") or d.get("hours") or ""
        ctx[f"D_{i}"] = d.get("D") or d.get("days") or ""
        ctx[f"UNITS_{i}"] = d.get("units") or ""
    # also expose `devices` list for templates that loop
    ctx["devices"] = devices
    ctx["device_count"] = len(devices)

    # Slabs (consumer copy)
    for i, s in enumerate(slabs, start=1):
        ctx[f"SLAB_{i}_UNITS"]  = s.get("yearly_units") or s.get("monthly_units") or ""
        ctx[f"SLAB_{i}_RATE"]   = s.get("rate") or ""
        ctx[f"SLAB_{i}_AMOUNT"] = _money(s.get("amount"))
    ctx["slabs"] = slabs

    # ---- Office identity (from system_config) — for PVVNL provisional notice
    ctx.update({
        "OFFICE_PHONE":     cfg.get("office_phone", ""),
        "OFFICE_EMAIL":     cfg.get("office_email", ""),
        "OFFICE_DIV_NO":    cfg.get("office_division_no", "") or cons.get("div_code", ""),
        "OFFICE_NAME_EN":   cfg.get("office_name_en", "Executive Engineer"),
        "OFFICE_NAME_HI":   cfg.get("office_name_hi", "अधिशासी अभियन्ता"),
        "OFFICE_DEPT_EN":   cfg.get("office_dept_en", "Electricity Distribution Division"),
        "OFFICE_DEPT_HI":   cfg.get("office_dept_hi", "विद्युत वितरण"),
        "OFFICE_LOC_EN":    cfg.get("office_location_en", ""),
        "OFFICE_LOC_HI":    cfg.get("office_location_hi", ""),
        "PATRANK_LETTER_CODE": cfg.get("patrank_letter_code", "वि0वि0ख0प्र0/शा0 एसैस्मेन्ट वि ."),
        "HEARING_OFFICER_ADDRESS_HI":
            cfg.get("hearing_officer_address_hi", ""),
    })

    # ---- Dispatch tracking fields (पत्रांक header)
    ctx.update({
        "DISPATCH_NUMBER":        case.get("dispatch_number") or extra.get("dispatch_number") or "",
        "DISPATCH_DATE":          case.get("dispatch_date")   or extra.get("dispatch_date")   or today_iso,
        "CHECKING_REPORT_NUMBER": case.get("checking_report_number")
                                  or extra.get("checking_report_number") or "",
        "HEARING_DATE":           case.get("hearing_date") or extra.get("hearing_date") or "",
        "HEARING_TIME":           case.get("hearing_time") or extra.get("hearing_time") or "",
    })

    # Allow caller-supplied overrides last
    ctx.update(extra)
    return ctx


def _load_office_config() -> dict:
    """Load the office_* config keys from system_config table."""
    try:
        rows = fetch_all(
            "SELECT config_key, config_value FROM system_config "
            "WHERE config_key LIKE 'office_%' OR config_key='patrank_letter_code' "
            "   OR config_key='hearing_officer_address_hi'"
        )
        return {r["config_key"]: r["config_value"] for r in rows}
    except Exception:  # noqa: BLE001
        return {}


# ===================================================================
# Legacy «FIELD» → {{ FIELD }} template migration
# ===================================================================
_LEGACY_RE = re.compile(r"«\s*([^»\s]+)\s*»")


def migrate_legacy_template_in_place(src: Path) -> Path:
    """
    Rewrite a .docx so every ``«FIELD»`` becomes ``{{ FIELD }}``.
    A backup is saved alongside as ``<name>.legacy.docx``.
    """
    if not src.exists():
        raise FileNotFoundError(src)
    backup = src.with_suffix(".legacy.docx")
    if not backup.exists():
        shutil.copy(src, backup)

    doc = Document(str(src))
    def _rewrite(text: str) -> str:
        return _LEGACY_RE.sub(lambda m: "{{ " + m.group(1).strip() + " }}", text)

    for p in doc.paragraphs:
        for run in p.runs:
            run.text = _rewrite(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.text = _rewrite(run.text)
    doc.save(str(src))
    log.info("Migrated legacy placeholders in %s (backup: %s)",
             src.name, backup.name)
    return backup


# ===================================================================
# Auto-generated documents (fallback when no template exists)
# ===================================================================
def _add_kv_table(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(items):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = str(v)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _autogen_provisional_consumer(out: Path, ctx: dict) -> None:
    doc = Document()
    _heading(doc, "अनन्तिम मूल्यांकन सूचना (Provisional Notice — Consumer Copy)", 0)
    doc.add_paragraph(f"दिनांक / Date: {ctx['Date']}    Case ID: {ctx['disno']}")

    _heading(doc, "उपभोक्ता विवरण / Consumer Details", 2)
    _add_kv_table(doc, [
        ("Account No.",   ctx["ACCOUNT_ID"]),
        ("Name",          ctx["NAME"]),
        ("Father / Husband", ctx["father_nane"]),
        ("Village",       ctx["VILLAGE"]),
        ("Post / Pin",    f"{ctx['post']} / {ctx['pin_code']}"),
        ("Mobile",        ctx["MOBILE_NO"]),
        ("Category",      ctx["CATEGORY"]),
        ("Section",       ctx["SECTION"]),
        ("Inspection Date", ctx["dis_date"]),
        ("Connected Load", ctx["CONNECTED_LOAD"]),
        ("Sub Station",   ctx["SUB_SUBSTATION"]),
        ("J.E.",          ctx["JE_NAME"]),
    ])

    _heading(doc, "उपकरण विवरण / Devices", 2)
    devices = ctx.get("devices") or []
    if devices:
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["S.No", "Device", "Load (W)", "Hours/day",
                               "Days", "Units"]):
            hdr[i].text = h
        for i, d in enumerate(devices, start=1):
            r = tbl.add_row().cells
            r[0].text = str(i)
            r[1].text = str(d.get("name", ""))
            r[2].text = str(d.get("L") or d.get("load") or "")
            r[3].text = str(d.get("H") or d.get("hours") or "")
            r[4].text = str(d.get("D") or d.get("days") or "")
            r[5].text = str(d.get("units") or "")

    _heading(doc, "मूल्यांकन / Assessment Breakdown", 2)
    _add_kv_table(doc, [
        ("Total Units (calculated)", ctx["TOTAL_CALCULATED_UNITS"]),
        ("Final Fixed Charges",      ctx["FINAL_FIXED_CHARGES"]),
        ("Final Energy Charges",     ctx["FINAL_ENERGY_CHARGES"]),
        ("Electricity Duty",         ctx["FINAL_ED_CHARGES"]),
        ("Grand Total",              ctx["ASSESMENT_TOTAL"]),
    ])

    p = doc.add_paragraph()
    p.add_run(
        "\nभुगतान की समय सीमा / Payment deadline: 7 days. "
        "आपत्ति / Appeal window: 15 days."
    )

    doc.save(str(out))


def _autogen_provisional_office(out: Path, ctx: dict) -> None:
    doc = Document()
    _heading(doc, "अनन्तिम मूल्यांकन (Office Copy — LFHD)", 0)
    doc.add_paragraph(f"Case ID: {ctx['disno']}    Date: {ctx['Date']}")
    _add_kv_table(doc, [
        ("Account No.",     ctx["ACCOUNT_ID"]),
        ("Name",            ctx["NAME"]),
        ("Section",         ctx["SECTION"]),
        ("Inspection Date", ctx["dis_date"]),
    ])

    _heading(doc, "LFHD Calculation (device names hidden)", 2)
    devices = ctx.get("devices") or []
    if devices:
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["#", "L", "F", "H", "D", "Units"]):
            tbl.rows[0].cells[i].text = h
        for i, d in enumerate(devices, start=1):
            r = tbl.add_row().cells
            r[0].text = str(i)
            r[1].text = str(d.get("L") or d.get("load") or "")
            r[2].text = str(d.get("F") or d.get("factor") or "")
            r[3].text = str(d.get("H") or d.get("hours") or "")
            r[4].text = str(d.get("D") or d.get("days") or "")
            r[5].text = str(d.get("units") or "")

    _heading(doc, "Summary", 2)
    _add_kv_table(doc, [
        ("Total Calculated Units", ctx["TOTAL_CALCULATED_UNITS"]),
        ("Summary — Fixed",        ctx["SUMMARY_FIXED"]),
        ("Summary — Energy",       ctx["SUMMARY_ENERGY"]),
        ("Summary — ED",           ctx["SUMMARY_ED"]),
        ("Summary — TOTAL",        ctx["SUMMARY_TOTAL"]),
    ])
    doc.save(str(out))


def _autogen_section3(out: Path, ctx: dict) -> None:
    doc = Document()
    _heading(doc, "Section 3 — Demand Notice", 0)
    doc.add_paragraph(f"Notice No.: {ctx.get('sec3_no','')}    "
                      f"Date: {ctx.get('sec3_date','')}")
    doc.add_paragraph(
        f"To,\n{ctx['NAME']}\nS/o {ctx['father_nane']}\n"
        f"Village {ctx['VILLAGE']}, Post {ctx['post']}, "
        f"Pin {ctx['pin_code']}\nMobile: {ctx['MOBILE_NO']}\n"
    )
    doc.add_paragraph(
        f"In reference to inspection dated {ctx['dis_date']} "
        f"(Case No. {ctx['disno']}), provisional assessment of "
        f"₹{ctx['ASSESMENT_TOTAL']} has been raised. "
        f"Pay within 30 days. Administrative charges ₹"
        f"{config.ADMIN_FEE_SECTION_3:.2f} apply."
    )
    _add_kv_table(doc, [
        ("Assessment Amount", ctx["sec3_amaunt"]),
        ("Admin Charges",     f"{config.ADMIN_FEE_SECTION_3:.2f}"),
        ("Total Payable",     ctx["total_sec3"]),
    ])
    doc.save(str(out))


def _autogen_section5(out: Path, ctx: dict) -> None:
    doc = Document()
    _heading(doc, "Section 5 — Revenue Recovery (RC)", 0)
    doc.add_paragraph(f"RC No.: {ctx.get('rc_number','')}   "
                      f"Letter: {ctx.get('letter_number','')}   "
                      f"Date: {ctx.get('current_date','')}")
    doc.add_paragraph(
        f"To,\nThe Tehsildar / Collector,\n{ctx['tehsil']}, "
        f"{ctx['district']}\n"
    )
    doc.add_paragraph(
        f"Recovery is requested against {ctx['consumer_name']} "
        f"S/o {ctx['consumer_father']}, R/o {ctx['village']} "
        f"({ctx['post_office']}). Outstanding ₹{ctx['outstanding_amount']}. "
        f"Reference: Checking report {ctx.get('checking_report_number','')} "
        f"dated {ctx.get('checking_date','')}, "
        f"Demand notice {ctx.get('demand_notice_number','')} "
        f"dated {ctx.get('demand_notice_date','')}."
    )
    doc.save(str(out))


def _autogen_thanedari(out: Path, ctx: dict) -> None:
    doc = Document()
    _heading(doc, "Thanedari Copy — Police / SE / Advocate", 0)
    doc.add_paragraph(f"Case ID: {ctx['disno']}    "
                      f"FIR: {ctx['FIR_NUMBER']}    Section: {ctx['SECTION']}")
    _add_kv_table(doc, [
        ("Consumer",       ctx["NAME"]),
        ("Father",         ctx["father_nane"]),
        ("Village",        ctx["VILLAGE"]),
        ("Mobile",         ctx["MOBILE_NO"]),
        ("Inspection",     ctx["dis_date"]),
        ("J.E.",           ctx["JE_NAME"]),
        ("Sub Station",    ctx["SUB_SUBSTATION"]),
        ("Assessment",     ctx["ASSESMENT_TOTAL"]),
        ("Compounding",    ctx["COMPOUNDING_AMOUNT"]),
    ])
    doc.save(str(out))


def _autogen_envelope(out: Path, ctx: dict) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(11)
    sec.page_width  = Cm(22)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("To,\n").bold = True
    addr = (
        f"{ctx['NAME']}\n"
        f"S/o {ctx['father_nane']}\n"
        f"Village {ctx['VILLAGE']}\n"
        f"Post {ctx['post']}, Pin {ctx['pin_code']}\n"
        f"Mob: {ctx['MOBILE_NO']}"
    )
    r = p.add_run(addr)
    r.font.size = Pt(14)
    doc.save(str(out))


def _autogen_deposit_slip(out: Path, ctx: dict) -> None:
    doc = Document()
    _heading(doc, "जमा पर्ची / Deposit Slip", 0)
    _add_kv_table(doc, [
        ("Date",        ctx["Date"]),
        ("Case ID",     ctx["disno"]),
        ("Account No.", ctx["ACCOUNT_ID"]),
        ("Division",    ctx["Div_no"]),
        ("Name",        ctx["NAME"]),
        ("Father",      ctx["father_nane"]),
        ("Village",     ctx["VILLAGE"]),
    ])
    _heading(doc, "Amounts", 2)
    _add_kv_table(doc, [
        ("Assessment Amount",   ctx["ASSESMENT_TOTAL"]),
        ("Compounding Charges", ctx["COMPOUNDING_AMOUNT"]),
    ])

    p = doc.add_paragraph()
    note = p.add_run(
        "\nयह धनराशि जमा करने के उपरान्त प्राप्त रसीद की छायाप्रति खण्ड "
        "कार्यालय के राजस्व निर्धारण पटल पर अवश्य जमा करा दें जिससे "
        "भविष्य की कार्यवाही जैसे पोर्टल पर अपलोड करना पत्र निर्गत् करना "
        "आदि हो सके।"
    )
    note.bold = True
    note.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    doc.save(str(out))


def _autogen_compounding_order(out: Path, ctx: dict) -> None:
    doc = Document()
    _heading(doc, "धारा 152 — Compounding Order", 0)
    _add_kv_table(doc, [
        ("Case ID", ctx["disno"]),
        ("Name",    ctx["NAME"]),
        ("Account", ctx["ACCOUNT_ID"]),
        ("Section", ctx["SECTION"]),
        ("Date",    ctx["Date"]),
    ])
    p = doc.add_paragraph()
    p.add_run("\n" + (ctx.get("COMPOUNDING_JUSTIFICATION") or ""))
    p.add_run(f"\n\nCompounding Amount: ₹{ctx['COMPOUNDING_AMOUNT']}")
    doc.save(str(out))


def _autogen_noc(out: Path, ctx: dict) -> None:
    doc = Document()
    _heading(doc, "No Objection Certificate (NOC)", 0)
    doc.add_paragraph(f"Case ID: {ctx['disno']}    Date: {ctx['Date']}")
    doc.add_paragraph(
        f"This is to certify that {ctx['NAME']} S/o {ctx['father_nane']}, "
        f"R/o {ctx['VILLAGE']}, holding account number "
        f"{ctx['ACCOUNT_ID']}, has cleared all dues against case "
        f"{ctx['disno']}. Accordingly NO OBJECTION is conveyed for "
        f"further action."
    )
    doc.save(str(out))


# =====================================================================
# PVVNL / UPPCL Provisional Notice — EXACT FORMAT
# =====================================================================
def _fmt_dmy(iso_date: str | None) -> str:
    """Convert ISO yyyy-mm-dd to dd/mm/yyyy."""
    if not iso_date:
        return ""
    try:
        dt = datetime.strptime(str(iso_date)[:10], "%Y-%m-%d")
        return dt.strftime("%d/%m/%Y")
    except (ValueError, TypeError):
        return str(iso_date)


def _build_lfhd_block(devices: list, assessment: dict) -> list[str]:
    """
    Build the LFHD calculation block exactly as in the PVVNL notice format:
        LFHD .963*1.00*18*365 =6327
        Total Unit=6327KWH
        EC 5.5*2*6327=69597.00
        FIXD 1*110 *2*12 =2640.00
        ED 7.5%@ =5418.00
        Tatal assessment=77655.00
    """
    lines = []
    total_units = 0.0

    for d in devices:
        load_w = safe_float(d.get("L") or d.get("load"))
        load_kw = round(load_w / 1000.0, 4) if load_w >= 1 else load_w
        # If load is already given as KW (e.g. 0.963), don't divide
        if load_w < 1 and load_w > 0:
            load_kw = load_w
        f_val = safe_float(d.get("F") or d.get("factor") or 1)
        h_val = safe_float(d.get("H") or d.get("hours"))
        d_val = safe_float(d.get("D") or d.get("days"))
        units = safe_float(d.get("units"))
        if units == 0 and load_kw > 0:
            units = round(load_kw * f_val * h_val * d_val, 2)
        total_units += units

        # Format: LFHD .963*1.00*18*365 =6327
        load_str = f"{load_kw:g}" if load_kw < 1 else f"{load_kw:g}"
        f_str = f"{f_val:.2f}" if f_val != int(f_val) else f"{int(f_val)}.00"
        lines.append(
            f"LFHD {load_str}*{f_str}*{int(h_val) if h_val == int(h_val) else h_val}*"
            f"{int(d_val) if d_val == int(d_val) else d_val} ={int(units) if units == int(units) else units:g}"
        )

    lines.append(f"Total Unit={int(total_units) if total_units == int(total_units) else round(total_units,2)}KWH")
    lines.append("")  # blank line

    a = assessment or {}
    energy = a.get("energy_charges") or {}
    fixed = a.get("fixed_charges") or {}
    ed = a.get("electricity_duty") or {}
    multiplier = safe_float(a.get("multiplier", 2))
    slabs = energy.get("slabs") or []

    # EC lines (one per slab) — format: rate*multiplier*units=amount
    for s in slabs:
        rate = safe_float(s.get("rate"))
        units = safe_float(s.get("yearly_units") or s.get("monthly_units"))
        amt = safe_float(s.get("amount"))
        rate_str = f"{rate:g}"
        if units > 0:
            lines.append(f"EC {rate_str}*{multiplier:g}*{int(units) if units == int(units) else units:g}={amt:.2f}")
        else:
            lines.append(f"{rate_str}*{multiplier:g}*=0.00")

    # FIXD line: load*fixed_rate*multiplier*months = base*multiplier
    cload_kw = safe_float(fixed.get("connected_load_kw"))
    fixed_rate = safe_float(fixed.get("fixed_rate"))
    months = safe_float(fixed.get("months"))
    fixed_final = safe_float(fixed.get("final"))
    months_int = int(round(months)) if months else 12
    cload_str = f"{cload_kw:g}" if cload_kw else "0"
    lines.append(
        f"FIXD {cload_str}*{fixed_rate:g} *{multiplier:g}*{months_int} ={fixed_final:.2f}"
    )

    # ED line
    ed_pct = safe_float(ed.get("ed_percent"))
    ed_amt = safe_float(ed.get("amount"))
    lines.append(f"ED {ed_pct:g}%@ ={ed_amt:.2f}")

    # Total
    grand = safe_float(a.get("grand_total"))
    lines.append(f"Tatal assessment={grand:.2f}")

    return lines


def _autogen_provisional_pvvnl(out: Path, ctx: dict) -> None:
    """
    Generate provisional notice in EXACT PVVNL/UPPCL format matching
    the user-supplied template (Shamli Division image).
    """
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    sec = doc.sections[0]
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

    # ============ HEADER ROW (3 cols: contact / logo / office) ============
    hdr_tbl = doc.add_table(rows=1, cols=3)
    hdr_tbl.autofit = False
    hdr_left = hdr_tbl.rows[0].cells[0]
    hdr_mid  = hdr_tbl.rows[0].cells[1]
    hdr_right = hdr_tbl.rows[0].cells[2]

    # Left: phone + email
    hl = hdr_left.paragraphs[0]
    hl.add_run(f"☎ {ctx.get('OFFICE_PHONE','')}\n").font.size = Pt(10)
    hl.add_run(f"✉ {ctx.get('OFFICE_EMAIL','')}").font.size = Pt(10)

    # Middle: placeholder for logo (centered text instead since no image)
    hm = hdr_mid.paragraphs[0]
    hm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    seal_run = hm.add_run("⚡")
    seal_run.font.size = Pt(28)
    seal_run.bold = True

    # Right: office name in Hindi/English
    hr = hdr_right.paragraphs[0]
    hr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = hr.add_run("कार्यालय\n")
    r1.font.size = Pt(10)
    r2 = hr.add_run(f"Office of the\n")
    r2.font.size = Pt(10)
    r3 = hr.add_run(f"{ctx.get('OFFICE_NAME_HI','अधिशासी अभियन्ता')}\n")
    r3.font.size = Pt(11); r3.bold = True
    r4 = hr.add_run(f"{ctx.get('OFFICE_NAME_EN','Executive Engineer')}\n")
    r4.font.size = Pt(10); r4.bold = True
    r5 = hr.add_run(f"{ctx.get('OFFICE_DEPT_HI','विद्युत वितरण')}\n")
    r5.font.size = Pt(10)
    r6 = hr.add_run(
        f"{ctx.get('OFFICE_DEPT_EN','Electricity Distribution Division')} "
        f"–{ctx.get('OFFICE_LOC_EN','')}\n"
    )
    r6.font.size = Pt(10)
    r7 = hr.add_run(f"{ctx.get('OFFICE_LOC_HI','')}")
    r7.font.size = Pt(10); r7.bold = True

    # Div No.
    p = doc.add_paragraph()
    p.add_run(f"Div No.- {ctx.get('OFFICE_DIV_NO', ctx.get('Div_no',''))}").bold = True

    # ============ पत्रांक line ============
    p = doc.add_paragraph()
    pr = p.add_run(
        f"पत्रांक    {ctx.get('DISPATCH_NUMBER','')}      "
        f"/{ctx.get('PATRANK_LETTER_CODE','')} {ctx.get('CHECKING_REPORT_NUMBER','')}                  "
        f"दिनांक: {_fmt_dmy(ctx.get('DISPATCH_DATE'))}"
    )
    pr.font.size = Pt(10); pr.bold = True

    doc.add_paragraph("")

    # ============ Heading box: "प्रस्तावित राजस्व निर्धारण की नोटिस" ============
    box_tbl = doc.add_table(rows=1, cols=3)
    box_tbl.autofit = False
    box_tbl.rows[0].cells[0].text = ""
    bm = box_tbl.rows[0].cells[1]
    bp = bm.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run("प्रस्तावित राजस्व निर्धारण की नोटिस")
    br.bold = True; br.font.size = Pt(14)
    # Right cell — पंजीकृत डाक stamp
    bright = box_tbl.rows[0].cells[2]
    bright_p = bright.paragraphs[0]
    bright_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    br2 = bright_p.add_run("पंजीकृत डाक")
    br2.bold = True; br2.font.size = Pt(11)

    # ============ विषय (Subject) ============
    insp_dmy = _fmt_dmy(ctx.get('dis_date'))
    p = doc.add_paragraph()
    sb = p.add_run("विषय – ")
    sb.bold = True
    p.add_run(
        f"दिनांक {insp_dmy} के निरीक्षण में पायी गयी अनियमितता के सन्दर्भ में "
        f"विद्युत अधिनियम, 2003 की धारा {ctx.get('SECTION','')} के अधीन "
        f"इलेक्ट्रिसिटी सप्लाई कोड–2005 के आलोक में प्रस्तावित राजस्व निर्धारण की नोटिस।"
    )

    # ============ महोदय,
    doc.add_paragraph("महोदय,")

    p = doc.add_paragraph()
    p.add_run(
        f"           आप अवगत ही होंगे कि दिनांक {insp_dmy} को आपके परिसर स्थापित विद्युत संयोजन का "
        f"निरीक्षण किया गया था जिसका विवरण निम्नवत है।"
    )

    # ============ Consumer details block (2 columns) ============
    cd_tbl = doc.add_table(rows=4, cols=2)
    cd_tbl.autofit = False
    rows_data = [
        (f"उपयोगकर्ता-श्री {ctx.get('USER_NAME','')} पुत्र {ctx.get('USERS_FATHER','')}",
         f"संयोजन संख्या -{ctx.get('ACCOUNT_ID','')}"),
        (f"संयोजन/परिसर स्वामी का नाम –श्री/श्रीमती-{ctx.get('NAME','')}",
         f"स्वीकृतभार – {ctx.get('CONNECTION_LOAD','')}"),
        (f"संयोजन/परिसर स्वामी के पिता/पति का नाम-{ctx.get('father_nane','')}",
         f"विद्या-{ctx.get('CATEGORY','')}"),
        (f"पता-{ctx.get('VILLAGE','')} पोस्ट-{ctx.get('post','')}",
         f"मो0नं0– {ctx.get('MOBILE_NO','')}"),
    ]
    for i, (l, r) in enumerate(rows_data):
        cd_tbl.rows[i].cells[0].text = l
        cd_tbl.rows[i].cells[1].text = r
    pin_p = doc.add_paragraph()
    pin_p.add_run(f"पिन कोड-{ctx.get('pin_code','')}")

    # ============ Body legal text ============
    section = ctx.get('SECTION', '135')
    p = doc.add_paragraph()
    p.add_run(
        f"           उक्त निरीक्षण में विद्युत अधिनियम, 2003 की धारा {section}/138/अन्य "
        f"के अंतर्गत दोषी पाए गए। जांचोपरांत आप द्वारा मौके पर शमन शुल्क जमा नहीं कराया गया है। "
        f"अतः आपके विरुद्ध प्रथमिकी भी दर्ज कराया जा चुका है।"
    )

    p = doc.add_paragraph()
    grand_str = ctx.get('ASSESMENT_TOTAL', '0.00')
    p.add_run(
        f"           अग्रेत्तर आपको अवगत कराना है कि इलेक्ट्रिसिटी सप्लाई कोड–2005 के आलोक में "
        f"उक्त अनियमितता पर प्रस्तावित राजस्व निर्धारण रू0 {grand_str} है। "
        f"आप उक्त प्रस्तावित राजस्व निर्धारण से संतुष्ट हैं, तो 7 दिनों के अंदर रू0 {grand_str} "
        f"जमा करना सुनिश्चित करे। यदि उक्त निर्धारण से संतुष्ट नहीं हैं, तो अपना दृष्टिकोण पूर्ण "
        f"आकार के कागज पर दो प्रतियों में साफ-साफ लिखकर सम्यक रूप से हस्ताक्षरित, साक्ष्य सामग्री "
        f"सहित 15 दिनों के अन्दर आप या किसी अन्य अधिकृत करते हुए, राजस्व निर्धारण अधिकारी "
        f"(खण्ड कार्यालय में ) को प्रस्तुत करे, अन्यथा प्रस्तावित राजस्व निर्धारण पर अग्रेत्तर "
        f"नियमानुसार कार्यवाही करते हुए राजस्व निर्धारण किया जायेगा। आपकी सुविधा हेतु आपका पक्ष "
        f"प्रस्तुत करने हेतु निम्नांकित तिथि निर्धारित की जाती है।"
    )

    # ============ Hearing officer + date/time ============
    p = doc.add_paragraph()
    p.add_run("सुनवाई हेतु अधिकृत अधिकारी का पदनाम एवं पता:- ").bold = True
    p.add_run(f"  {ctx.get('HEARING_OFFICER_ADDRESS_HI','')}")

    hearing_dmy = _fmt_dmy(ctx.get('HEARING_DATE')) or "—"
    hearing_time = ctx.get('HEARING_TIME') or "—"
    p = doc.add_paragraph()
    p.add_run("निर्धारित तिथि– ").bold = True
    p.add_run(f"{hearing_dmy}                समय– {hearing_time}")

    p = doc.add_paragraph()
    p.add_run(
        "             यदि आपके द्वारा निर्धारित तिथि या 15 दिनों तक अपना प्रत्यावेदन/पक्ष प्रस्तुत नहीं "
        "किया जाता है तो प्रकरण का निस्तारण नियमानुसार आपके पक्ष की अनुपस्थिति में किया जायेगा। "
        "जिसका सम्पूर्ण उत्तरदायित्व आपका होगा।"
    )

    p = doc.add_paragraph()
    nr = p.add_run("नोट –: ")
    nr.bold = True
    p.add_run("राजस्व निर्धारण के अतिरिक्त शमन शुल्क नियमानुसार देय होगा ।")

    # ============ LFHD calculation block ============
    p = doc.add_paragraph()
    p.add_run("प्रस्तावित राजस्व निर्धारण की गणना निम्न प्रकार हैः–").bold = True

    # Build LFHD lines
    devices = ctx.get("devices") or []
    # Need assessment dict — reconstruct from ctx fields if possible
    assessment = {
        "energy_charges": {
            "slabs": ctx.get("slabs") or [],
            "final": safe_float(ctx.get("FINAL_ENERGY_CHARGES", "0").replace(",", "")),
        },
        "fixed_charges": {
            "connected_load_kw": safe_float(ctx.get("CONNECTED_LOAD") or 0),
            "fixed_rate":  safe_float(str(ctx.get("FIXED_RATE","0")).replace(",", "")),
            "months":      safe_float(ctx.get("MONTHS") or 12),
            "final":       safe_float(str(ctx.get("FINAL_FIXED_CHARGES","0")).replace(",", "")),
        },
        "electricity_duty": {
            "ed_percent": safe_float(ctx.get("ED_RATE_PERCENT") or 0),
            "amount":     safe_float(str(ctx.get("FINAL_ED_CHARGES","0")).replace(",", "")),
        },
        "multiplier": 2,
        "grand_total": safe_float(str(ctx.get("ASSESMENT_TOTAL","0")).replace(",", "")),
    }
    lines = _build_lfhd_block(devices, assessment)
    for line in lines:
        lp = doc.add_paragraph()
        run = lp.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)

    # ============ Signature block ============
    doc.add_paragraph("")
    sig_tbl = doc.add_table(rows=2, cols=2)
    sig_tbl.rows[0].cells[0].text = ""
    sig_right = sig_tbl.rows[0].cells[1]
    sr_p = sig_right.paragraphs[0]
    sr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sr_run = sr_p.add_run(f"{ctx.get('OFFICE_NAME_HI','अधिशासी अभियन्ता')}/निर्धारण अधिकारी")
    sr_run.bold = True
    sr2 = sig_tbl.rows[1].cells[1].paragraphs[0]
    sr2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sr2.add_run("O/C").bold = True

    doc.save(str(out))


_AUTOGEN_HANDLERS = {
    "provisional_consumer": _autogen_provisional_consumer,
    "provisional_office":   _autogen_provisional_office,
    "provisional_pvvnl":    _autogen_provisional_pvvnl,
    "section3":             _autogen_section3,
    "section5":             _autogen_section5,
    "thanedari":            _autogen_thanedari,
    "envelope":             _autogen_envelope,
    "deposit_slip":         _autogen_deposit_slip,
    "compounding_order":    _autogen_compounding_order,
    "noc":                  _autogen_noc,
}


# ===================================================================
# Public API
# ===================================================================
def generate(case_id: str, kind: str,
             extra: dict | None = None,
             user: str = "system") -> dict:
    """Generate a document for a case. Returns metadata + file_path."""
    if kind not in VALID_KINDS:
        raise ValueError(f"Unknown kind '{kind}'. Valid: {VALID_KINDS}")

    case = fetch_one("SELECT * FROM raid_cases WHERE case_id=?", (case_id,))
    if not case:
        raise LookupError(f"Case {case_id} not found")
    consumer = (fetch_one("SELECT * FROM consumers WHERE id=?",
                          (case["consumer_id"],))
                if case["consumer_id"] else None)
    case["devices"] = from_json_str(case.get("devices_json")) or []
    assessment = from_json_str(case.get("assessment_json")) or {}

    ctx = build_context(case, consumer, assessment, extra=extra)

    # Tag compounding justification if present
    if not ctx.get("COMPOUNDING_JUSTIFICATION"):
        from .compounding import calculate_compounding
        load_kw = safe_float(case.get("connected_load_kw"))
        if load_kw > 0:
            comp = calculate_compounding({
                "load_kw": load_kw,
                "category": (consumer or {}).get("category"),
                "section":  case.get("section"),
            })
            if comp.get("ok"):
                ctx["COMPOUNDING_JUSTIFICATION"] = comp["justification_hi"]
                ctx["COMPOUNDING_AMOUNT"] = _money(comp["compounding_amount"])

    # Output path: docs/<case_id>/<kind>_<timestamp>.docx
    out_dir = config.DOCS_DIR / case_id
    out_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_file = out_dir / f"{kind}_{timestamp}.docx"

    template_path = config.TEMPLATES_DIR / KIND_TO_TEMPLATE[kind]
    used_template = False

    if template_path.exists():
        try:
            tpl = DocxTemplate(str(template_path))
            tpl.render(ctx)
            tpl.save(str(out_file))
            used_template = True
        except Exception as e:  # noqa: BLE001
            log.exception("Template render failed for %s; falling back to autogen",
                          template_path.name)
            # Fall through to autogen
            used_template = False

    if not used_template:
        handler = _AUTOGEN_HANDLERS.get(kind)
        if not handler:
            raise NotImplementedError(f"No autogen handler for kind '{kind}'")
        handler(out_file, ctx)

    file_size = out_file.stat().st_size
    cur = execute(
        """INSERT INTO documents
              (case_id, document_type, document_name, file_path,
               file_size, mime_type, uploaded_by)
           VALUES (?,?,?,?,?,?,?)""",
        (case_id, kind, out_file.name, str(out_file), file_size,
         "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
         user),
    )

    return {
        "id": cur.lastrowid,
        "case_id": case_id,
        "kind": kind,
        "used_template": used_template,
        "template_file": str(template_path) if used_template else None,
        "file_path": str(out_file),
        "file_name": out_file.name,
        "file_size": file_size,
        "context_keys": sorted(ctx.keys()),
    }
