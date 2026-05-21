#!/usr/bin/env python3
"""
Generate default Word (.docx) templates for the Raid Management System.

These templates use docxtpl-compatible ``{{ FIELD }}`` placeholders so they
can be rendered by ``backend.services.doc_generator`` without modification.

Run:
    python scripts/generate_default_templates.py

The script writes 9 .docx files into ``templates/`` (overwriting any
existing file with the same name). Officers can then open any of these
in Microsoft Word and edit text/styling while keeping the placeholders.

The placeholders follow the project spec's mail-merge field catalog
(see .kiro/steering/project-spec.md "Mail Merge Field Catalog").
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
TEMPLATES = ROOT / "templates"
TEMPLATES.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------- helpers
def _heading(doc: Document, text: str, level: int = 0,
             align: int = WD_ALIGN_PARAGRAPH.CENTER) -> None:
    h = doc.add_heading(text, level=level)
    h.alignment = align


def _kv(doc: Document, items: list[tuple[str, str]]) -> None:
    """Two-column key-value table."""
    tbl = doc.add_table(rows=len(items), cols=2)
    tbl.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(items):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
    for row in tbl.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _para(doc: Document, text: str, *, bold: bool = False, size: int = 11,
          color: tuple[int, int, int] | None = None,
          align: int = WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def _save(doc: Document, name: str) -> Path:
    out = TEMPLATES / name
    doc.save(str(out))
    print(f"[OK] {out.relative_to(ROOT)}")
    return out


# ============================================================ TEMPLATES

def tpl_provisional_consumer():
    doc = Document()
    _heading(doc, "अनन्तिम मूल्यांकन सूचना", 0)
    _heading(doc, "Provisional Notice — Consumer Copy", 1)
    _para(doc,
          "दिनांक / Date: {{ Date }}    "
          "Case No.: {{ disno }}    Division: {{ Div_no }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    _heading(doc, "उपभोक्ता विवरण / Consumer Details", 2,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    _kv(doc, [
        ("Account No.",        "{{ ACCOUNT_ID }}"),
        ("Online No.",         "{{ ONLINE_NO }}"),
        ("Name / नाम",          "{{ NAME }}"),
        ("Father / पिता",       "{{ father_nane }}"),
        ("Village / गाँव",       "{{ VILLAGE }}"),
        ("Post / Pin",         "{{ post }} / {{ pin_code }}"),
        ("Mobile",             "{{ MOBILE_NO }}"),
        ("Category",           "{{ CATEGORY }}"),
        ("Supply Type",        "{{ SUPPLY_TYPE }}"),
        ("Section",            "{{ SECTION }}"),
        ("Inspection Date",    "{{ dis_date }}"),
        ("Connected Load (kW)", "{{ CONNECTED_LOAD }}"),
        ("Sub-Substation",     "{{ SUB_SUBSTATION }}"),
        ("J.E.",               "{{ JE_NAME }}"),
        ("Checking Type",      "{{ CHECKING_TYPE }}"),
    ])

    _heading(doc, "उपकरण विवरण / Devices (LFHD)", 2,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    _para(doc,
          "{% for d in devices %}"
          "{{ loop.index }}. {{ d.name }} — "
          "Load {{ d.L or d.load }} W × Hours {{ d.H or d.hours }} × "
          "Days {{ d.D or d.days }} = {{ d.units }} units\n"
          "{% endfor %}")

    _heading(doc, "मूल्यांकन सारांश / Assessment Summary", 2,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    _kv(doc, [
        ("Total Units (calculated)", "{{ TOTAL_CALCULATED_UNITS }}"),
        ("Months",                   "{{ MONTHS }}"),
        ("Fixed Rate (₹/kW/mo)",     "{{ FIXED_RATE }}"),
        ("Final Fixed Charges",      "₹ {{ FINAL_FIXED_CHARGES }}"),
        ("Final Energy Charges",     "₹ {{ FINAL_ENERGY_CHARGES }}"),
        ("Electricity Duty ({{ ED_RATE_PERCENT }}%)",
                                     "₹ {{ FINAL_ED_CHARGES }}"),
        ("GRAND TOTAL",              "₹ {{ ASSESMENT_TOTAL }}"),
        ("Compounding (Sec 152)",    "₹ {{ COMPOUNDING_AMOUNT }}"),
    ])

    _para(doc,
          "\nभुगतान की समय सीमा 7 दिन है। आपत्ति की समय सीमा 15 दिन है। "
          "Payment due within 7 days; appeal window 15 days.",
          bold=True, color=(192, 0, 0))

    _para(doc, "\n\nहस्ताक्षर / Signature: ____________________     "
               "(अधिशासी अभियन्ता / Executive Engineer)")
    return _save(doc, "provisional_consumer.docx")


def tpl_provisional_office():
    doc = Document()
    _heading(doc, "अनन्तिम मूल्यांकन — कार्यालय प्रति", 0)
    _heading(doc, "Provisional Notice — Office Copy (LFHD focus)", 1)
    _para(doc,
          "Case No.: {{ disno }}    Division: {{ Div_no }}    "
          "Date: {{ Date }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    _kv(doc, [
        ("Account No.",     "{{ ACCOUNT_ID }}"),
        ("Online No.",      "{{ ONLINE_NO }}"),
        ("Name",            "{{ NAME }}"),
        ("Section",         "{{ SECTION }}"),
        ("Inspection Date", "{{ dis_date }}"),
        ("J.E.",            "{{ JE_NAME }}"),
        ("Sub-Substation",  "{{ SUB_SUBSTATION }}"),
    ])

    _heading(doc, "LFHD Calculation (device names suppressed)", 2,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    _para(doc,
          "{% for d in devices %}"
          "Item {{ loop.index }}:  L={{ d.L or d.load }}  "
          "F={{ d.F or d.factor }}  H={{ d.H or d.hours }}  "
          "D={{ d.D or d.days }}  →  {{ d.units }} units\n"
          "{% endfor %}")

    _heading(doc, "Office Summary", 2, align=WD_ALIGN_PARAGRAPH.LEFT)
    _kv(doc, [
        ("Total Calculated Units", "{{ TOTAL_CALCULATED_UNITS }}"),
        ("Summary — Fixed",   "₹ {{ SUMMARY_FIXED }}"),
        ("Summary — Energy",  "₹ {{ SUMMARY_ENERGY }}"),
        ("Summary — ED",      "₹ {{ SUMMARY_ED }}"),
        ("Summary — TOTAL",   "₹ {{ SUMMARY_TOTAL }}"),
    ])
    _para(doc, "\n\nप्रति / Copy: मुख्य अभियन्ता, अधिशासी अभियन्ता, "
               "उपखण्ड अभियन्ता, राजस्व निर्धारण पटल",
          size=10)
    return _save(doc, "provisional_office.docx")


def tpl_section3():
    doc = Document()
    _heading(doc, "धारा 3 के अंतर्गत मांग सूचना", 0)
    _heading(doc, "Section 3 — Demand Notice", 1)
    _para(doc,
          "Notice No.: {{ sec3_no }}    Date: {{ sec3_date }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    _para(doc,
          "\nसेवा में,\n{{ NAME }}\nपुत्र / पति  {{ father_nane }}\n"
          "ग्राम  {{ VILLAGE }}, पोस्ट  {{ post }}, पिन  {{ pin_code }}\n"
          "मोबाइल: {{ MOBILE_NO }}")

    _para(doc,
          "\nविषय: निरीक्षण दिनांक {{ dis_date }} (केस सं. {{ disno }}) "
          "के आधार पर अनन्तिम मूल्यांकन ₹ {{ ASSESMENT_TOTAL }} "
          "की वसूली हेतु धारा 3 के अंतर्गत मांग सूचना।")

    _para(doc,
          "\nमहोदय,\nउपरोक्त प्रकरण में आपके परिसर पर निरीक्षण दिनांक "
          "{{ dis_date }} को जे.ई. {{ JE_NAME }} द्वारा विद्युत चोरी / "
          "अनधिकृत उपयोग पाया गया। जिसके आधार पर अनन्तिम मूल्यांकन "
          "₹ {{ ASSESMENT_TOTAL }} किया गया है। प्रशासनिक शुल्क "
          "₹ 25.00 अतिरिक्त देय होगा।\n\nकृपया 30 दिन के भीतर "
          "पूर्ण भुगतान कर रसीद की प्रति इस कार्यालय में जमा करें, "
          "अन्यथा धारा 5 के अंतर्गत राजस्व वसूली की कार्यवाही प्रारम्भ "
          "की जाएगी।")

    _kv(doc, [
        ("Assessment Amount",   "₹ {{ sec3_amaunt }}"),
        ("Admin Charges",       "₹ 25.00"),
        ("TOTAL PAYABLE",       "₹ {{ total_sec3 }}"),
    ])
    _para(doc, "\n\nभवदीय,\nअधिशासी अभियन्ता\nखण्ड {{ Div_no }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    return _save(doc, "section3_notice.docx")


def tpl_section5():
    doc = Document()
    _heading(doc, "धारा 5 — राजस्व वसूली प्रमाण-पत्र", 0)
    _heading(doc, "Section 5 — Revenue Recovery Certificate (RRC)", 1)
    _para(doc,
          "RC No.: {{ rc_number }}    Letter No.: {{ letter_number }}    "
          "Date: {{ current_date }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    _para(doc,
          "\nसेवा में,\nश्रीमान तहसीलदार / जिलाधिकारी,\n"
          "तहसील  {{ tehsil }},  जनपद  {{ district }}\n")

    _para(doc,
          "\nविषय: उपभोक्ता {{ consumer_name }} पुत्र {{ consumer_father }} "
          "के विरुद्ध बकाया राशि ₹ {{ outstanding_amount }} की राजस्व वसूली "
          "हेतु प्रमाण-पत्र निर्गमन।")

    _para(doc,
          "\nमहोदय,\nउपरोक्त उपभोक्ता निवासी ग्राम {{ village }}, "
          "पोस्ट {{ post_office }} पर इस कार्यालय की निम्नलिखित "
          "कार्यवाही के बावजूद बकाया राशि का भुगतान नहीं किया गया है:")

    _kv(doc, [
        ("Checking Report",  "{{ checking_report_number }} dt {{ checking_date }}"),
        ("Demand Notice",    "{{ demand_notice_number }} dt {{ demand_notice_date }}"),
        ("Grid / SE Letter", "{{ grid_number }} dt {{ grid_date }}"),
        ("Outstanding",      "₹ {{ outstanding_amount }}"),
    ])

    _para(doc,
          "\nअतः अनुरोध है कि उक्त राशि भू-राजस्व की भांति वसूल कर "
          "इस कार्यालय को सूचित करें।")
    _para(doc, "\n\nभवदीय,\nअधिशासी अभियन्ता\nखण्ड {{ Div_no }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    return _save(doc, "section5_notice.docx")


def tpl_thanedari():
    doc = Document()
    _heading(doc, "थानेदारी प्रति", 0)
    _heading(doc, "Thanedari / Police Copy", 1)
    _para(doc,
          "Case No.: {{ disno }}    FIR: {{ FIR_NUMBER }}    "
          "Section: {{ SECTION }}    Date: {{ Date }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    _para(doc, "\nसेवा में,\nप्रभारी निरीक्षक,\nथाना ____________________")

    _para(doc,
          "\nविषय: विद्युत अधिनियम धारा {{ SECTION }} के अंतर्गत प्रकरण "
          "संख्या {{ disno }} में सहयोग एवं FIR दर्ज करवाने हेतु।")

    _kv(doc, [
        ("उपभोक्ता / Consumer",  "{{ NAME }}"),
        ("पुत्र / Father",        "{{ father_nane }}"),
        ("ग्राम / Village",       "{{ VILLAGE }}"),
        ("मोबाइल / Mobile",       "{{ MOBILE_NO }}"),
        ("Account",              "{{ ACCOUNT_ID }}"),
        ("निरीक्षण / Inspection", "{{ dis_date }}"),
        ("J.E.",                 "{{ JE_NAME }}"),
        ("Sub-Substation",       "{{ SUB_SUBSTATION }}"),
        ("मूल्यांकन / Assessment", "₹ {{ ASSESMENT_TOTAL }}"),
        ("Compounding",          "₹ {{ COMPOUNDING_AMOUNT }}"),
    ])

    _para(doc, "\nप्रति / Cc: अधीक्षण अभियन्ता, कॉर्पोरेट अधिवक्ता",
          size=10)
    _para(doc, "\n\nभवदीय,\nअधिशासी अभियन्ता\nखण्ड {{ Div_no }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    return _save(doc, "thanedari_copy.docx")


def tpl_envelope():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(11)
    sec.page_width = Cm(22)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("From:\nExecutive Engineer\nDivision  {{ Div_no }}\n"
                  "U.P. Power Corporation Ltd.")
    r.font.size = Pt(11)

    doc.add_paragraph("\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("To,\n")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("{{ NAME }}\n"
                  "S/o {{ father_nane }}\n"
                  "Village {{ VILLAGE }}\n"
                  "Post {{ post }} - {{ pin_code }}\n"
                  "Mob: {{ MOBILE_NO }}")
    r.font.size = Pt(14)
    return _save(doc, "envelope.docx")


def tpl_deposit_slip():
    doc = Document()
    _heading(doc, "जमा पर्ची", 0)
    _heading(doc, "Deposit Slip", 1)
    _para(doc,
          "Case No.: {{ disno }}    Division: {{ Div_no }}    "
          "Date: {{ Date }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    _kv(doc, [
        ("Account No.", "{{ ACCOUNT_ID }}"),
        ("Online No.",  "{{ ONLINE_NO }}"),
        ("Name",        "{{ NAME }}"),
        ("Father",      "{{ father_nane }}"),
        ("Village",     "{{ VILLAGE }}"),
        ("Mobile",      "{{ MOBILE_NO }}"),
        ("Section",     "{{ SECTION }}"),
        ("J.E.",        "{{ JE_NAME }}"),
    ])

    _heading(doc, "देय राशि / Amounts Due", 2,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    _kv(doc, [
        ("Assessment Amount",     "₹ {{ ASSESMENT_TOTAL }}"),
        ("Compounding Charges",   "₹ {{ COMPOUNDING_AMOUNT }}"),
    ])

    # Mandatory Hindi reminder (red, bold) — verbatim from project spec
    _para(doc,
          "\nयह धनराशि जमा करने के उपरान्त प्राप्त रसीद की छायाप्रति "
          "खण्ड कार्यालय के राजस्व निर्धारण पटल पर अवश्य जमा करा दें "
          "जिससे भविष्य की कार्यवाही जैसे पोर्टल पर अपलोड करना पत्र "
          "निर्गत् करना आदि हो सके।",
          bold=True, color=(192, 0, 0))

    _para(doc,
          "\n\nखण्ड कार्यालय की मुहर / Office Stamp:    "
          "       कैशियर हस्ताक्षर / Cashier Sign:",
          size=10)
    return _save(doc, "deposit_slip.docx")


def tpl_compounding_order():
    doc = Document()
    _heading(doc, "धारा 152 — Compounding आदेश", 0)
    _heading(doc, "Section 152 Compounding Order", 1)
    _para(doc,
          "Case No.: {{ disno }}    Date: {{ Date }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    _kv(doc, [
        ("Consumer",   "{{ NAME }}"),
        ("Father",     "{{ father_nane }}"),
        ("Village",    "{{ VILLAGE }}"),
        ("Account",    "{{ ACCOUNT_ID }}"),
        ("Section",    "{{ SECTION }}"),
        ("Connected Load", "{{ CONNECTED_LOAD }} kW"),
        ("Compounding Amount", "₹ {{ COMPOUNDING_AMOUNT }}"),
    ])

    _heading(doc, "औचित्य / Justification", 2,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    _para(doc, "{{ COMPOUNDING_JUSTIFICATION }}")

    _para(doc,
          "\nउपरोक्त आधार पर धारा 152 के अंतर्गत Compounding शुल्क "
          "₹ {{ COMPOUNDING_AMOUNT }} निर्धारित किया जाता है।",
          bold=True)

    _para(doc, "\n\nअधिशासी अभियन्ता\nखण्ड {{ Div_no }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    return _save(doc, "compounding_order.docx")


def tpl_noc():
    doc = Document()
    _heading(doc, "अनापत्ति प्रमाण-पत्र", 0)
    _heading(doc, "No Objection Certificate (NOC)", 1)
    _para(doc, "Case No.: {{ disno }}    Date: {{ Date }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    _para(doc,
          "\nप्रमाणित किया जाता है कि श्री / श्रीमती  {{ NAME }}  पुत्र / "
          "पति  {{ father_nane }}, निवासी ग्राम  {{ VILLAGE }}, "
          "पोस्ट  {{ post }}, खाता संख्या  {{ ACCOUNT_ID }} "
          "के विरुद्ध केस संख्या  {{ disno }}  में निर्धारित सम्पूर्ण "
          "राशि ₹ {{ ASSESMENT_TOTAL }} एवं Compounding ₹ "
          "{{ COMPOUNDING_AMOUNT }} का भुगतान कर दिया गया है।\n\n"
          "अतः इस कार्यालय की ओर से कोई आपत्ति शेष नहीं है। This NOC is "
          "issued for further legal/portal action.")

    _para(doc, "\n\nभवदीय,\nअधिशासी अभियन्ता\nखण्ड {{ Div_no }}",
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    return _save(doc, "noc.docx")


# ============================================================ main
def main() -> None:
    print(f"Writing default templates to: {TEMPLATES}")
    tpl_provisional_consumer()
    tpl_provisional_office()
    tpl_section3()
    tpl_section5()
    tpl_thanedari()
    tpl_envelope()
    tpl_deposit_slip()
    tpl_compounding_order()
    tpl_noc()
    print(f"\nDone. {len(list(TEMPLATES.glob('*.docx')))} templates ready.")


if __name__ == "__main__":
    main()
