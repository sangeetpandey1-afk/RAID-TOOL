#!/usr/bin/env python3
"""
Generate all 9 document template .docx files for the RAID-TOOL.

Each template uses Jinja2-style {{ FIELD }} placeholders compatible with
docxtpl. Hindi text is included directly (UTF-8). On the user's Windows
machine, the font can be changed to Krutidev manually if needed.

Run:  python scripts/generate_templates.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)


def _set_font(run, name="Kruti Dev 010", size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold


def _add_heading(doc, text, level=1, align=WD_ALIGN_PARAGRAPH.CENTER):
    h = doc.add_heading(text, level=level)
    h.alignment = align
    return h



def _add_table_row(table, label, value):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value


def _kv_table(doc, items):
    """Create a 2-column key-value table."""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in items:
        _add_table_row(table, k, v)
    return table


# ===================================================================
# 1. PROVISIONAL NOTICE — CONSUMER COPY
# ===================================================================
def create_provisional_consumer():
    doc = Document()
    _add_heading(doc, "अनन्तिम मूल्यांकन सूचना पत्र", 0)
    _add_heading(doc, "(Provisional Assessment Notice — Consumer Copy)", 2)

    doc.add_paragraph("कार्यालय — अवर अभियन्ता, विद्युत वितरण खण्ड")
    doc.add_paragraph("संख्या: {{ disno }}    दिनांक: {{ Date }}")
    doc.add_paragraph("")


    # Consumer details table
    _add_heading(doc, "उपभोक्ता विवरण / Consumer Details", 2)
    _kv_table(doc, [
        ("खण्ड / Division", "{{ Div_no }}"),
        ("ऑनलाइन नं. / Online No.", "{{ ONLINE_NO }}"),
        ("खाता संख्या / Account No.", "{{ ACCOUNT_ID }}"),
        ("उपभोक्ता नाम / Name", "{{ NAME }}"),
        ("पिता/पति नाम / Father/Husband", "{{ father_nane }}"),
        ("ग्राम / Village", "{{ VILLAGE }}"),
        ("डाकघर / Post", "{{ post }}"),
        ("पिन कोड / Pin", "{{ pin_code }}"),
        ("मोबाइल / Mobile", "{{ MOBILE_NO }}"),
        ("श्रेणी / Category", "{{ CATEGORY }}"),
        ("भार / Connected Load", "{{ CONNECTION_LOAD }}"),
        ("आपूर्ति प्रकार / Supply Type", "{{ SUPPLY_TYPE }}"),
        ("उप-केन्द्र / Sub Station", "{{ SUB_SUBSTATION }}"),
        ("अवर अभियन्ता / J.E.", "{{ JE_NAME }}"),
        ("निरीक्षण दिनांक / Inspection Date", "{{ dis_date }}"),
        ("धारा / Section", "{{ SECTION }}"),
        ("जाँच प्रकार / Checking Type", "{{ CHECKING_TYPE }}"),
    ])
    doc.add_paragraph("")


    # Devices table with names (consumer copy shows device names)
    _add_heading(doc, "उपकरण विवरण / Device Details", 2)
    dev_table = doc.add_table(rows=1, cols=6)
    dev_table.style = "Table Grid"
    headers = ["क्र.सं.", "उपकरण नाम / Device", "भार (W) / Load",
               "घण्टे / Hours", "दिन / Days", "यूनिट / Units"]
    for i, h in enumerate(headers):
        dev_table.rows[0].cells[i].text = h

    # Up to 15 device rows with placeholders
    for n in range(1, 16):
        row = dev_table.add_row()
        row.cells[0].text = str(n)
        row.cells[1].text = "{{ " + f"DEVICE_{n}_NAME" + " }}"
        row.cells[2].text = "{{ " + f"DEVICE_{n}_LOAD" + " }}"
        row.cells[3].text = "{{ " + f"H_{n}" + " }}"
        row.cells[4].text = "{{ " + f"D_{n}" + " }}"
        row.cells[5].text = "{{ " + f"DEVICE_{n}_UNITS" + " }}"
    doc.add_paragraph("")


    # Assessment breakdown
    _add_heading(doc, "मूल्यांकन विवरण / Assessment Breakdown", 2)
    _kv_table(doc, [
        ("कुल गणित यूनिट / Total Calculated Units", "{{ TOTAL_CALCULATED_UNITS }}"),
        ("स्थिर दर / Fixed Rate", "{{ FIXED_RATE }}"),
        ("माह / Months", "{{ MONTHS }}"),
        ("स्थिर शुल्क राशि / Fixed Amount", "{{ FIXED_AMOUNT }}"),
    ])
    doc.add_paragraph("")

    # Slab-wise energy charges
    _add_heading(doc, "स्लैब-वार ऊर्जा शुल्क / Slab-wise Energy Charges", 2)
    slab_table = doc.add_table(rows=1, cols=3)
    slab_table.style = "Table Grid"
    slab_table.rows[0].cells[0].text = "स्लैब यूनिट / Slab Units"
    slab_table.rows[0].cells[1].text = "दर / Rate"
    slab_table.rows[0].cells[2].text = "राशि / Amount"
    for n in range(1, 5):
        row = slab_table.add_row()
        row.cells[0].text = "{{ " + f"SLAB_{n}_UNITS" + " }}"
        row.cells[1].text = "{{ " + f"SLAB_{n}_RATE" + " }}"
        row.cells[2].text = "{{ " + f"SLAB_{n}_AMOUNT" + " }}"
    doc.add_paragraph("")


    # Final totals
    _add_heading(doc, "कुल देय राशि / Total Payable", 2)
    _kv_table(doc, [
        ("अन्तिम स्थिर शुल्क / Final Fixed Charges", "{{ FINAL_FIXED_CHARGES }}"),
        ("अन्तिम ऊर्जा शुल्क / Final Energy Charges", "{{ FINAL_ENERGY_CHARGES }}"),
        ("विद्युत शुल्क (ED) % / ED Rate", "{{ ED_RATE_PERCENT }}"),
        ("विद्युत शुल्क राशि / ED Amount", "{{ FINAL_ED_CHARGES }}"),
        ("कुल मूल्यांकन / Grand Total", "{{ ASSESMENT_TOTAL }}"),
    ])

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("भुगतान की अन्तिम तिथि: ").bold = True
    p.add_run("निरीक्षण दिनांक से 7 दिन के भीतर")
    p = doc.add_paragraph()
    p.add_run("आपत्ति/अपील: ").bold = True
    p.add_run("15 दिन के भीतर प्रस्तुत करें")

    doc.add_paragraph("\n\nहस्ताक्षर अवर अभियन्ता")

    out = TEMPLATES_DIR / "provisional_consumer.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")



# ===================================================================
# 2. PROVISIONAL NOTICE — OFFICE COPY (LFHD, no device names)
# ===================================================================
def create_provisional_office():
    doc = Document()
    _add_heading(doc, "अनन्तिम मूल्यांकन — कार्यालय प्रति", 0)
    _add_heading(doc, "(Provisional Assessment — Office Copy / LFHD)", 2)

    doc.add_paragraph("Case ID: {{ disno }}    Date: {{ Date }}")
    doc.add_paragraph("")

    _kv_table(doc, [
        ("खाता संख्या / Account No.", "{{ ACCOUNT_ID }}"),
        ("नाम / Name", "{{ NAME }}"),
        ("पिता नाम / Father", "{{ father_nane }}"),
        ("धारा / Section", "{{ SECTION }}"),
        ("निरीक्षण दिनांक / Inspection Date", "{{ dis_date }}"),
        ("खण्ड / Division", "{{ Div_no }}"),
    ])
    doc.add_paragraph("")

    # LFHD table (device names hidden)
    _add_heading(doc, "LFHD Calculation", 2)
    lfhd_table = doc.add_table(rows=1, cols=6)
    lfhd_table.style = "Table Grid"
    headers = ["#", "L (Load)", "F (Factor)", "H (Hours)", "D (Days)", "Units"]
    for i, h in enumerate(headers):
        lfhd_table.rows[0].cells[i].text = h


    for n in range(1, 16):
        row = lfhd_table.add_row()
        row.cells[0].text = str(n)
        row.cells[1].text = "{{ " + f"L_{n}" + " }}"
        row.cells[2].text = "{{ " + f"F_{n}" + " }}"
        row.cells[3].text = "{{ " + f"H_{n}" + " }}"
        row.cells[4].text = "{{ " + f"D_{n}" + " }}"
        row.cells[5].text = "{{ " + f"UNITS_{n}" + " }}"
    doc.add_paragraph("")

    # Summary
    _add_heading(doc, "Summary / सारांश", 2)
    _kv_table(doc, [
        ("Total Calculated Units", "{{ TOTAL_CALCULATED_UNITS }}"),
        ("Fixed Charges", "{{ SUMMARY_FIXED }}"),
        ("Energy Charges", "{{ SUMMARY_ENERGY }}"),
        ("Electricity Duty (ED)", "{{ SUMMARY_ED }}"),
        ("GRAND TOTAL", "{{ SUMMARY_TOTAL }}"),
    ])

    doc.add_paragraph("\n\nहस्ताक्षर अवर अभियन्ता")

    out = TEMPLATES_DIR / "provisional_office.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")



# ===================================================================
# 3. SECTION 3 NOTICE (30-day demand + ₹25 admin fee)
# ===================================================================
def create_section3():
    doc = Document()
    _add_heading(doc, "धारा 3 — माँग सूचना पत्र", 0)
    _add_heading(doc, "(Section 3 — Demand Notice)", 2)

    doc.add_paragraph("कार्यालय — अवर अभियन्ता, विद्युत वितरण खण्ड")
    doc.add_paragraph("सूचना संख्या: {{ sec3_no }}    दिनांक: {{ sec3_date }}")
    doc.add_paragraph("खण्ड: {{ div_no }}")
    doc.add_paragraph("")

    doc.add_paragraph("सेवा में,")
    doc.add_paragraph("श्री/श्रीमती {{ NAME }}")
    doc.add_paragraph("पुत्र/पति श्री {{ father_nane }}")
    doc.add_paragraph("ग्राम — {{ VILLAGE }}")
    doc.add_paragraph("डाकघर — {{ post }}, पिन — {{ pin_code }}")
    doc.add_paragraph("मोबाइल — {{ MOBILE_NO }}")
    doc.add_paragraph("")

    doc.add_paragraph("विषय: विद्युत चोरी के सम्बन्ध में माँग सूचना पत्र")
    doc.add_paragraph("")


    p = doc.add_paragraph()
    p.add_run("महोदय/महोदया,").bold = True
    doc.add_paragraph(
        "दिनांक {{ dis_date }} को आपके परिसर का निरीक्षण किया गया "
        "(Case No. {{ disno }}) जिसमें विद्युत चोरी/अनधिकृत उपयोग पाया गया। "
        "धारा {{ SECTION }} के अन्तर्गत अनन्तिम मूल्यांकन ₹{{ sec3_amaunt }} "
        "निर्धारित किया गया है।"
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "आपको सूचित किया जाता है कि उक्त राशि 30 दिन के भीतर जमा करें। "
        "विफल रहने पर धारा 5 के अन्तर्गत राजस्व वसूली की कार्यवाही की जाएगी।"
    )
    doc.add_paragraph("")

    _add_heading(doc, "देय राशि विवरण / Amount Details", 2)
    _kv_table(doc, [
        ("मूल्यांकन राशि / Assessment Amount", "₹{{ sec3_amaunt }}"),
        ("प्रशासनिक शुल्क / Admin Charges", "₹25.00"),
        ("कुल देय / Total Payable", "₹{{ total_sec3 }}"),
        ("खाता संख्या / Account No.", "{{ conaction_no }}"),
        ("कनेक्शन संख्या / Connection No.", "{{ conaction_no }}"),
    ])

    doc.add_paragraph("\n\nदिनांक: {{ Date }}")
    doc.add_paragraph("हस्ताक्षर अवर अभियन्ता")

    out = TEMPLATES_DIR / "section3_notice.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")



# ===================================================================
# 4. SECTION 5 NOTICE (Revenue Recovery via Tehsildar/Collector)
# ===================================================================
def create_section5():
    doc = Document()
    _add_heading(doc, "धारा 5 — राजस्व वसूली पत्र", 0)
    _add_heading(doc, "(Section 5 — Revenue Recovery Certificate)", 2)

    doc.add_paragraph("कार्यालय — अवर अभियन्ता, विद्युत वितरण खण्ड")
    doc.add_paragraph(
        "RC संख्या: {{ rc_number }}    "
        "पत्र संख्या: {{ letter_number }}    "
        "दिनांक: {{ current_date }}"
    )
    doc.add_paragraph("")

    doc.add_paragraph("सेवा में,")
    doc.add_paragraph("तहसीलदार / जिला कलेक्टर")
    doc.add_paragraph("तहसील — {{ tehsil }}")
    doc.add_paragraph("जनपद — {{ district }}")
    doc.add_paragraph("")

    doc.add_paragraph("विषय: विद्युत देय राशि की राजस्व वसूली हेतु")
    doc.add_paragraph("")


    doc.add_paragraph("महोदय,")
    doc.add_paragraph(
        "निम्नलिखित उपभोक्ता के विरुद्ध विद्युत अधिनियम 2003 की धारा {{ SECTION }} "
        "के अन्तर्गत जाँच रिपोर्ट संख्या {{ checking_report_number }} "
        "दिनांक {{ checking_date }} द्वारा मूल्यांकन किया गया था। "
        "माँग सूचना संख्या {{ demand_notice_number }} दिनांक {{ demand_notice_date }} "
        "द्वारा भुगतान हेतु सूचित किया गया, परन्तु भुगतान प्राप्त नहीं हुआ।"
    )
    doc.add_paragraph("")

    _add_heading(doc, "उपभोक्ता विवरण / Consumer Details", 2)
    _kv_table(doc, [
        ("नाम / Name", "{{ consumer_name }}"),
        ("पिता/पति / Father", "{{ consumer_father }}"),
        ("ग्राम / Village", "{{ village }}"),
        ("डाकघर / Post Office", "{{ post_office }}"),
        ("तहसील / Tehsil", "{{ tehsil }}"),
        ("जनपद / District", "{{ district }}"),
        ("खाता संख्या / Account No.", "{{ ACCOUNT_ID }}"),
        ("बकाया राशि / Outstanding Amount", "₹{{ outstanding_amount }}"),
        ("ग्रिड संख्या / Grid No.", "{{ grid_number }}"),
        ("ग्रिड दिनांक / Grid Date", "{{ grid_date }}"),
    ])

    doc.add_paragraph(
        "\nअतः अनुरोध है कि उक्त बकाया राशि ₹{{ outstanding_amount }} "
        "की वसूली भू-राजस्व की भाँति कराने की कृपा करें।"
    )
    doc.add_paragraph("\n\nहस्ताक्षर अवर अभियन्ता")

    out = TEMPLATES_DIR / "section5_notice.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")



# ===================================================================
# 5. THANEDARI COPY (Police / SE Office / Corporate Advocate)
# ===================================================================
def create_thanedari():
    doc = Document()
    _add_heading(doc, "थानेदारी प्रति", 0)
    _add_heading(doc, "(Thanedari Copy — Police / SE / Advocate)", 2)

    doc.add_paragraph("कार्यालय — अवर अभियन्ता, विद्युत वितरण खण्ड")
    doc.add_paragraph("Case ID: {{ disno }}    दिनांक: {{ Date }}")
    doc.add_paragraph("")

    doc.add_paragraph("सेवा में,")
    doc.add_paragraph("थाना प्रभारी / अधीक्षण अभियन्ता / विधिक प्रतिनिधि")
    doc.add_paragraph("")

    doc.add_paragraph(
        "विषय: विद्युत चोरी प्रकरण में FIR/कार्यवाही हेतु"
    )
    doc.add_paragraph("")

    _add_heading(doc, "प्रकरण विवरण / Case Details", 2)
    _kv_table(doc, [
        ("नाम / Name", "{{ NAME }}"),
        ("पिता नाम / Father", "{{ father_nane }}"),
        ("ग्राम / Village", "{{ VILLAGE }}"),
        ("मोबाइल / Mobile", "{{ MOBILE_NO }}"),
        ("खाता संख्या / Account No.", "{{ ACCOUNT_ID }}"),
        ("निरीक्षण दिनांक / Inspection Date", "{{ dis_date }}"),
        ("धारा / Section", "{{ SECTION }}"),
        ("FIR संख्या / FIR Number", "{{ FIR_NUMBER }}"),
        ("अवर अभियन्ता / J.E.", "{{ JE_NAME }}"),
        ("उप-केन्द्र / Sub Station", "{{ SUB_SUBSTATION }}"),
        ("जाँच प्रकार / Checking Type", "{{ CHECKING_TYPE }}"),
        ("मूल्यांकन राशि / Assessment", "₹{{ ASSESMENT_TOTAL }}"),
        ("Compounding राशि", "₹{{ COMPOUNDING_AMOUNT }}"),
    ])

    doc.add_paragraph(
        "\nउक्त उपभोक्ता के विरुद्ध विद्युत अधिनियम 2003 की धारा {{ SECTION }} "
        "के अन्तर्गत कार्यवाही की जानी है। कृपया आवश्यक सहयोग प्रदान करें।"
    )
    doc.add_paragraph("\n\nहस्ताक्षर अवर अभियन्ता")

    out = TEMPLATES_DIR / "thanedari_copy.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")



# ===================================================================
# 6. ENVELOPE (printer-optimized address)
# ===================================================================
def create_envelope():
    doc = Document()
    # Envelope dimensions
    section = doc.sections[0]
    section.page_height = Cm(11)
    section.page_width = Cm(22)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.5)

    # From (left-aligned, smaller)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("प्रेषक / From:")
    r.bold = True
    r.font.size = Pt(10)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run(
        "अवर अभियन्ता\n"
        "विद्युत वितरण खण्ड\n"
        "{{ Div_no }}"
    )
    r2.font.size = Pt(10)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # To (right-aligned, larger)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p3.add_run("सेवा में / To,")
    r3.bold = True
    r3.font.size = Pt(12)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    addr = (
        "{{ NAME }}\n"
        "पुत्र/पति श्री {{ father_nane }}\n"
        "ग्राम — {{ VILLAGE }}\n"
        "निकट — {{ LANDMARK }}\n"
        "डाकघर — {{ post }}, पिन — {{ pin_code }}\n"
        "मोबाइल — {{ MOBILE_NO }}"
    )
    r4 = p4.add_run(addr)
    r4.font.size = Pt(14)

    out = TEMPLATES_DIR / "envelope.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")



# ===================================================================
# 7. DEPOSIT SLIP (with mandatory reminder text)
# ===================================================================
def create_deposit_slip():
    doc = Document()
    _add_heading(doc, "जमा पर्ची / Deposit Slip", 0)

    doc.add_paragraph("दिनांक: {{ Date }}    Case ID: {{ disno }}")
    doc.add_paragraph("")

    _add_heading(doc, "उपभोक्ता विवरण", 2)
    _kv_table(doc, [
        ("खाता संख्या / Account No.", "{{ ACCOUNT_ID }}"),
        ("खण्ड / Division", "{{ Div_no }}"),
        ("नाम / Name", "{{ NAME }}"),
        ("पिता/पति / Father", "{{ father_nane }}"),
        ("ग्राम / Village", "{{ VILLAGE }}"),
        ("मोबाइल / Mobile", "{{ MOBILE_NO }}"),
        ("धारा / Section", "{{ SECTION }}"),
    ])
    doc.add_paragraph("")

    _add_heading(doc, "जमा राशि विवरण / Amount Details", 2)
    _kv_table(doc, [
        ("मूल्यांकन राशि / Assessment Amount", "₹{{ ASSESMENT_TOTAL }}"),
        ("Compounding शुल्क / Compounding Charges", "₹{{ COMPOUNDING_AMOUNT }}"),
        ("कुल जमा / Total Deposit", "₹{{ ASSESMENT_TOTAL }}"),
    ])
    doc.add_paragraph("")


    # Mandatory reminder (bold, red)
    p = doc.add_paragraph()
    p.add_run("महत्वपूर्ण सूचना / Important Notice:").bold = True
    doc.add_paragraph("")

    reminder = doc.add_paragraph()
    r = reminder.add_run(
        "यह धनराशि जमा करने के उपरान्त प्राप्त रसीद की छायाप्रति खण्ड "
        "कार्यालय के राजस्व निर्धारण पटल पर अवश्य जमा करा दें जिससे "
        "भविष्य की कार्यवाही जैसे पोर्टल पर अपलोड करना पत्र निर्गत् करना "
        "आदि हो सके।"
    )
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph("\n\nजमाकर्ता हस्ताक्षर:  ____________")
    doc.add_paragraph("कार्यालय हस्ताक्षर:  ____________")
    doc.add_paragraph("दिनांक: {{ Date }}")

    out = TEMPLATES_DIR / "deposit_slip.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")



# ===================================================================
# 8. NOC (No Objection Certificate — on full payment)
# ===================================================================
def create_noc():
    doc = Document()
    _add_heading(doc, "अनापत्ति प्रमाण पत्र", 0)
    _add_heading(doc, "(No Objection Certificate — NOC)", 2)

    doc.add_paragraph("कार्यालय — अवर अभियन्ता, विद्युत वितरण खण्ड")
    doc.add_paragraph("Case ID: {{ disno }}    दिनांक: {{ Date }}")
    doc.add_paragraph("")

    doc.add_paragraph(
        "प्रमाणित किया जाता है कि श्री/श्रीमती {{ NAME }} "
        "पुत्र/पति श्री {{ father_nane }}, "
        "निवासी ग्राम {{ VILLAGE }}, "
        "खाता संख्या {{ ACCOUNT_ID }} ने "
        "प्रकरण संख्या {{ disno }} के विरुद्ध समस्त देय राशि जमा कर दी है।"
    )
    doc.add_paragraph("")

    _add_heading(doc, "प्रकरण विवरण", 2)
    _kv_table(doc, [
        ("खाता संख्या / Account No.", "{{ ACCOUNT_ID }}"),
        ("नाम / Name", "{{ NAME }}"),
        ("पिता/पति / Father", "{{ father_nane }}"),
        ("ग्राम / Village", "{{ VILLAGE }}"),
        ("धारा / Section", "{{ SECTION }}"),
        ("निरीक्षण दिनांक / Inspection Date", "{{ dis_date }}"),
        ("मूल्यांकन राशि / Assessment", "₹{{ ASSESMENT_TOTAL }}"),
        ("Compounding राशि", "₹{{ COMPOUNDING_AMOUNT }}"),
    ])
    doc.add_paragraph("")


    doc.add_paragraph(
        "अतः उक्त उपभोक्ता के विरुद्ध इस कार्यालय की ओर से कोई आपत्ति नहीं है। "
        "भविष्य में विद्युत आपूर्ति/अन्य सेवाओं हेतु यह NOC मान्य रहेगा।"
    )

    doc.add_paragraph("\n\nहस्ताक्षर अवर अभियन्ता")
    doc.add_paragraph("दिनांक: {{ Date }}")
    doc.add_paragraph("मुहर / Seal")

    out = TEMPLATES_DIR / "noc.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")


# ===================================================================
# 9. COMPOUNDING ORDER (Section 152)
# ===================================================================
def create_compounding_order():
    doc = Document()
    _add_heading(doc, "धारा 152 — शमन आदेश", 0)
    _add_heading(doc, "(Section 152 — Compounding Order)", 2)

    doc.add_paragraph("कार्यालय — अवर अभियन्ता, विद्युत वितरण खण्ड")
    doc.add_paragraph("Case ID: {{ disno }}    दिनांक: {{ Date }}")
    doc.add_paragraph("")


    _add_heading(doc, "उपभोक्ता विवरण", 2)
    _kv_table(doc, [
        ("खाता संख्या / Account No.", "{{ ACCOUNT_ID }}"),
        ("नाम / Name", "{{ NAME }}"),
        ("पिता/पति / Father", "{{ father_nane }}"),
        ("ग्राम / Village", "{{ VILLAGE }}"),
        ("खण्ड / Division", "{{ Div_no }}"),
        ("धारा / Section", "{{ SECTION }}"),
        ("निरीक्षण दिनांक / Inspection Date", "{{ dis_date }}"),
        ("भार (KW) / Connected Load", "{{ CONNECTION_LOAD }}"),
        ("श्रेणी / Category", "{{ CATEGORY }}"),
    ])
    doc.add_paragraph("")

    _add_heading(doc, "Compounding गणना / Calculation", 2)
    doc.add_paragraph(
        "विद्युत अधिनियम 2003 की धारा 152 के अन्तर्गत LT उपभोक्ताओं के "
        "मामलों में Compounding charges \"per KW or part thereof\" basis पर "
        "लगती हैं।"
    )
    doc.add_paragraph("")

    # Justification text placeholder
    p = doc.add_paragraph()
    r = p.add_run("{{ COMPOUNDING_JUSTIFICATION }}")
    r.bold = True
    doc.add_paragraph("")

    _kv_table(doc, [
        ("Compounding राशि / Amount", "₹{{ COMPOUNDING_AMOUNT }}"),
    ])

    doc.add_paragraph(
        "\nउक्त राशि 7 दिन के भीतर जमा करें अन्यथा FIR दर्ज कराई जाएगी।"
    )
    doc.add_paragraph("\n\nहस्ताक्षर अवर अभियन्ता")
    doc.add_paragraph("दिनांक: {{ Date }}")

    out = TEMPLATES_DIR / "compounding_order.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")



# ===================================================================
# MAIN — Generate all templates
# ===================================================================
if __name__ == "__main__":
    print("Generating document templates in:", TEMPLATES_DIR)
    print("=" * 50)

    create_provisional_consumer()
    create_provisional_office()
    create_section3()
    create_section5()
    create_thanedari()
    create_envelope()
    create_deposit_slip()
    create_noc()
    create_compounding_order()

    print("=" * 50)
    print("Done! All 9 templates created successfully.")
    print(f"Location: {TEMPLATES_DIR}")
